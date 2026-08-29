"""Strict DOCX text replacement while retaining source structure and media."""

from __future__ import annotations

import io
import logging
from hashlib import sha256
from itertools import accumulate
from typing import Any

from ..exceptions import DependencyUnavailableError, RenderingError
from ..schemas import (
    DocumentModel,
    LayoutStatus,
    ReconstructionType,
    RegionType,
    TranslationStatus,
)
from ..utils.text_utils import normalize_text

LOGGER = logging.getLogger(__name__)


def replace_paragraph_text_preserving_runs(paragraph: object, text: str) -> None:
    """Distribute translation across existing text runs without touching media runs.

    Assigning ``run.text`` replaces that run's XML children. Empty runs often own an
    inline image, drawing, field, or bookmark, so they must never be cleared merely
    because neighboring text is translated.
    """
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    active = [run for run in runs if run.text]
    if not active:
        runs[0].text = text
        return
    source_lengths = [max(1, len(run.text)) for run in active]
    total_source = sum(source_lengths)
    source_text = "".join(run.text for run in active)
    words = text.split()
    if not words:
        for run in active:
            run.text = ""
        return
    # Retain intentional hard-line structure when translation has no explicit
    # breaks. Paragraph spacing, tabs, indents, alignment, and styles remain on
    # their original paragraph/run objects.
    line_breaks = source_text.count("\n")
    if line_breaks and "\n" not in text and len(words) > line_breaks:
        line_count = line_breaks + 1
        split_at = {
            max(1, round(len(words) * line / line_count))
            for line in range(1, line_count)
        }
        expanded: list[str] = []
        for index, word in enumerate(words, start=1):
            expanded.append(word)
            if index in split_at:
                expanded.append("\n")
        words = expanded
    cumulative = list(accumulate(source_lengths))
    buckets = [""] * len(active)
    character_position = 0
    for word in words:
        is_break = word == "\n"
        target = min(
            len(active) - 1,
            next(
                (
                    index
                    for index, boundary in enumerate(cumulative)
                    if character_position / max(1, len(text)) <= boundary / total_source
                ),
                len(active) - 1,
            ),
        )
        if is_break:
            buckets[target] = buckets[target].rstrip() + "\n"
            character_position += 1
            continue
        separator = "" if not buckets[target] or buckets[target].endswith("\n") else " "
        buckets[target] = f"{buckets[target]}{separator}{word}"
        character_position += len(word) + 1
    for run in active:
        run.text = ""
    for run, value in zip(active, buckets, strict=True):
        run.text = value + (" " if value and run is not active[-1] else "")


def _run_contains_non_text_content(run: object) -> bool:
    element = getattr(run, "_r", None)
    if element is None:
        return False
    try:
        return bool(element.xpath(".//w:drawing | .//w:object | .//w:pict"))
    except Exception:
        return False


def _safe_replacement(block: object) -> tuple[bool, str]:
    """Return whether a translated block may alter the primary DOCX."""
    if block.translation_status != TranslationStatus.TRANSLATED:
        return False, "translation_not_completed"
    if not block.source_validated:
        return False, "source_not_validated"
    if block.detected_language not in {"pa", "hi"}:
        return False, "not_validated_punjabi_or_hindi"
    if not (block.english_translation or "").strip():
        return False, "empty_translation"
    if not block.source_reference or block.source_reference.startswith("embedded:"):
        return False, "no_safe_native_text_reference"
    if block.region_type in {
        RegionType.STAMP_SEAL,
        RegionType.SIGNATURE,
        RegionType.GRAPHICAL_CONTENT,
    }:
        return False, "critical_graphic_region"
    if block.reconstruction_type == ReconstructionType.UNREADABLE:
        return False, "unreadable_source"
    if block.metadata.get("preserve_region_as_image"):
        return False, "preserve_region_as_image"
    return True, "eligible"


def _preserve_original(block: object, reason: str) -> None:
    """Record a terminal, auditable no-op for a native DOCX text block."""
    block.output_bbox = None
    block.layout_status = LayoutStatus.SKIPPED
    block.metadata.update(
        {
            "render_mode": "preserve_original",
            "replacement_applied": False,
            "preserved_original": True,
            "preservation_reason": reason,
            "primary_output_replacement_reason": reason,
        }
    )


def _resolve_paragraph(document: object, reference: str) -> object | None:
    parts = reference.split(":")
    try:
        if parts[:2] == ["body", "p"]:
            return document.paragraphs[int(parts[2])]
        if parts[0] == "table":
            table_index, row_index, cell_index, paragraph_index = (
                int(parts[1]),
                int(parts[2]),
                int(parts[3]),
                int(parts[5]),
            )
            return document.tables[table_index].rows[row_index].cells[cell_index].paragraphs[paragraph_index]
        if parts[0] == "section":
            section = document.sections[int(parts[1])]
            region = section.header if parts[2] == "header" else section.footer
            return region.paragraphs[int(parts[4])]
    except (IndexError, ValueError):
        LOGGER.warning("DOCX source reference could not be resolved: %s", reference)
    return None


def _image_hashes(document: object) -> tuple[str, ...]:
    hashes: list[str] = []
    for part in document.part.package.parts:
        if str(getattr(part, "content_type", "")).startswith("image/"):
            hashes.append(sha256(bytes(part.blob)).hexdigest())
    return tuple(sorted(hashes))


def _structure_signature(document: object) -> dict[str, Any]:
    """Capture invariants that primary reconstruction is not allowed to change."""
    section_geometry = []
    header_footer_shapes = []
    for section in document.sections:
        section_geometry.append(
            (
                int(section.page_width or 0),
                int(section.page_height or 0),
                str(section.orientation),
                int(section.top_margin or 0),
                int(section.bottom_margin or 0),
                int(section.left_margin or 0),
                int(section.right_margin or 0),
            )
        )
        header_footer_shapes.append(
            (
                len(section.header.paragraphs),
                len(section.header.tables),
                len(section.footer.paragraphs),
                len(section.footer.tables),
            )
        )
    table_shape = tuple(
        (len(table.rows), tuple(len(row.cells) for row in table.rows))
        for table in document.tables
    )
    return {
        "sections": len(document.sections),
        "section_geometry": tuple(section_geometry),
        "body_paragraphs": len(document.paragraphs),
        "body_tables": table_shape,
        "inline_shapes": len(document.inline_shapes),
        "header_footer_shapes": tuple(header_footer_shapes),
        "images": _image_hashes(document),
    }


def reconstruct_docx(document_model: DocumentModel) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise DependencyUnavailableError("python-docx is required for DOCX reconstruction") from exc
    try:
        document = Document(io.BytesIO(document_model.source_bytes))
        source_signature = _structure_signature(document)
        references: dict[str, list[object]] = {}
        for block in document_model.blocks:
            allowed, reason = _safe_replacement(block)
            _preserve_original(block, reason)
            if allowed and block.source_reference:
                references.setdefault(block.source_reference, []).append(block)
        for reference, blocks in references.items():
            # Multiple translations targeting one paragraph cannot be mapped back to
            # individual runs safely. Fail closed and retain the original paragraph.
            eligible = []
            for block in blocks:
                allowed, reason = _safe_replacement(block)
                block.metadata["primary_output_replacement_reason"] = reason
                if allowed:
                    eligible.append(block)
            if len(eligible) != 1:
                if len(eligible) > 1:
                    for block in eligible:
                        block.metadata["primary_output_replacement_reason"] = (
                            "ambiguous_multiple_blocks_for_native_paragraph"
                        )
                continue
            block = eligible[0]
            paragraph = _resolve_paragraph(document, reference)
            if paragraph is None:
                block.metadata["primary_output_replacement_reason"] = "source_reference_not_found"
                continue
            if any(
                run.text and _run_contains_non_text_content(run)
                for run in paragraph.runs
            ):
                block.metadata["primary_output_replacement_reason"] = (
                    "text_and_graphic_share_one_run"
                )
                continue
            expected = normalize_text(
                str(block.metadata.get("native_source_snapshot", block.source_text))
            )
            actual = normalize_text(paragraph.text)
            if expected != actual:
                block.metadata["primary_output_replacement_reason"] = (
                    "native_source_snapshot_mismatch"
                )
                continue
            mutable_text = normalize_text(
                "".join(run.text for run in paragraph.runs if run.text)
            )
            if mutable_text != actual:
                # Hyperlink/field text can be visible through paragraph.text while
                # not belonging to paragraph.runs. Appending replacement text would
                # duplicate it, so preserve the paragraph intact.
                block.metadata["primary_output_replacement_reason"] = (
                    "text_not_owned_by_mutable_runs"
                )
                continue
            replace_paragraph_text_preserving_runs(
                paragraph, block.english_translation or ""
            )
            block.output_bbox = block.source_bbox
            block.layout_status = LayoutStatus.FIT
            block.metadata.update(
                {
                    "render_mode": "native_structure_replacement",
                    "replacement_applied": True,
                    "preserved_original": False,
                    "primary_output_replacement_reason": "replaced_in_native_structure",
                }
            )
            source_length = max(1, len(expected))
            length_deviation = abs(len(block.english_translation or "") - source_length) / source_length
            block.metadata["format_fidelity_score"] = max(
                0.60, 1.0 - min(1.0, length_deviation) * 0.25
            )
        buffer = io.BytesIO()
        document.save(buffer)
        output = buffer.getvalue()
        rebuilt = Document(io.BytesIO(output))
        if _structure_signature(rebuilt) != source_signature:
            raise RenderingError(
                "DOCX structural validation failed; the primary output was not exported"
            )
        for block in document_model.blocks:
            if block.metadata.get("replacement_applied") is not True:
                continue
            paragraph = _resolve_paragraph(rebuilt, str(block.source_reference or ""))
            if paragraph is None or normalize_text(paragraph.text) != normalize_text(
                block.english_translation or ""
            ):
                raise RenderingError(
                    "DOCX replacement verification failed; the primary output was not exported"
                )
        document_model.metadata["applied_translation_replacement_count"] = sum(
            block.translation_status == TranslationStatus.TRANSLATED
            and block.metadata.get("replacement_applied") is True
            for block in document_model.blocks
        )
        document_model.metadata["docx_structure_validated"] = True
        document_model.metadata["docx_primary_output_has_appended_diagnostics"] = False
        for page in document_model.pages:
            scores = [
                float(block.metadata["format_fidelity_score"])
                for block in page.blocks
                if isinstance(block.metadata.get("format_fidelity_score"), (int, float))
            ]
            page.metadata["format_fidelity_score"] = (
                sum(scores) / len(scores) if scores else 1.0
            )
            page.metadata["format_fidelity_method"] = (
                "native DOCX structural invariants and run-preserving replacement"
            )
        return output
    except Exception as exc:
        if isinstance(exc, (DependencyUnavailableError, RenderingError)):
            raise
        raise RenderingError("DOCX reconstruction failed") from exc
