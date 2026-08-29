"""Model-free tests for fast PDF/DOCX document routing."""

from __future__ import annotations

from docx import Document
import pymupdf
from PIL import Image

from pretrained_page_ocr import docx_processor
import pretrained_page_ocr.document_processor as document_processor
from pretrained_page_ocr.document_processor import (
    _ExtractedLine,
    _line_result,
    process_fast_image,
    process_pdf,
)


def test_native_pdf_text_bypasses_ocr_and_creates_word_output(tmp_path) -> None:
    source_path = tmp_path / "native.pdf"
    document = pymupdf.open()
    page = document.new_page(width=420, height=600)
    page.insert_text((40, 70), "Existing English document text")
    document.save(source_path)
    document.close()

    payload = process_pdf(source_path, tmp_path / "output")

    assert payload["pages"][0]["route"] == "native_text"
    assert "Existing English document text" in payload["english_translation"]
    assert payload["models"]["gurmukhi"]["loaded"] is False
    assert payload["models"]["hindi"]["loaded"] is False
    assert (tmp_path / "output" / "translated_en.docx").is_file()


def test_short_native_gurmukhi_is_not_rejected_by_ocr_length_rules() -> None:
    extracted = _ExtractedLine(
        text="ਪੱਤਰ",
        bbox=(10, 20, 100, 45),
        confidence=0.99,
        provider="pymupdf_native_text",
    )

    result = _line_result(
        extracted,
        page_number=1,
        line_number=1,
        crop_file="",
    )

    assert result.accepted is True
    assert result.review_required is False
    assert result.script == "gurmukhi"


def test_pdf_with_native_text_and_page_image_uses_mixed_scan_route(
    monkeypatch, tmp_path
) -> None:
    image_path = tmp_path / "background.png"
    Image.new("RGB", (420, 600), "white").save(image_path)
    source_path = tmp_path / "mixed.pdf"
    document = pymupdf.open()
    page = document.new_page(width=420, height=600)
    page.insert_image(page.rect, filename=str(image_path))
    page.insert_text((35, 55), "Native selectable hospital form heading")
    document.save(source_path)
    document.close()
    monkeypatch.setattr(
        document_processor,
        "_tesseract_pdf_lines",
        lambda page: [
            _ExtractedLine(
                "Scanned handwritten area",
                (40, 300, 350, 340),
                0.92,
                "tesseract_pdf_fast",
            )
        ],
    )
    monkeypatch.setattr(
        document_processor,
        "_supplement_scanned_regions",
        lambda page, lines: sorted(lines, key=lambda item: item.bbox[1]),
    )

    payload = process_pdf(source_path, tmp_path / "output")

    assert payload["pages"][0]["route"] == "mixed_native_scanned"
    assert len(payload["lines"]) == 2


def test_native_docx_preserves_structure_and_replaces_validated_indic(
    monkeypatch, tmp_path
) -> None:
    source_path = tmp_path / "letter.docx"
    document = Document()
    document.add_paragraph("Existing English heading")
    document.add_paragraph("ਪੰਜਾਬੀ ਭਾਸ਼ਾ ਦੀ ਲਿਖਤ")
    document.save(source_path)

    def fake_translate(lines, models, progress_callback=None, num_beams=4):
        assert num_beams == 4
        for line in lines:
            if line.accepted and line.script in {"gurmukhi", "devanagari"}:
                line.english = "Punjabi language text"
                line.translation_status = "translated"
                line.review_required = False
        return 0.01

    monkeypatch.setattr(docx_processor, "translate_text", fake_translate)
    payload = docx_processor.process_docx(source_path, tmp_path / "output")
    rebuilt = Document(tmp_path / "output" / "translated_en.docx")

    assert [paragraph.text for paragraph in rebuilt.paragraphs] == [
        "Existing English heading",
        "Punjabi language text",
    ]
    assert payload["native_translation_replacement_count"] == 1
    assert payload["lines"][0]["accepted"] is True


def test_docx_split_keeps_numbers_out_of_source_translation() -> None:
    parts = docx_processor.split_translatable_parts("ਪੱਤਰ 2026 ਪੰਜਾਬੀ")

    assert "".join(part.text for part in parts) == "ਪੱਤਰ 2026 ਪੰਜਾਬੀ"
    assert any(part.script is None and "2026" in part.text for part in parts)


def test_fast_image_route_does_not_load_neural_handwriting_model(
    monkeypatch, tmp_path
) -> None:
    source_path = tmp_path / "scan.png"
    Image.new("RGB", (500, 700), "white").save(source_path)
    monkeypatch.setattr(
        document_processor,
        "_tesseract_pdf_lines",
        lambda page: [
            _ExtractedLine(
                "Existing English scan text",
                (30, 50, 430, 90),
                0.95,
                "tesseract_pdf_fast",
            )
        ],
    )
    monkeypatch.setattr(
        document_processor,
        "_supplement_scanned_regions",
        lambda page, lines: list(lines),
    )

    payload = process_fast_image(source_path, tmp_path / "output")

    assert payload["models"]["gurmukhi"]["loaded"] is False
    assert payload["models"]["hindi"]["loaded"] is False
    assert payload["lines"][0]["translation_status"] == "preserved_english"
    assert (tmp_path / "output" / "translated_en.docx").is_file()
