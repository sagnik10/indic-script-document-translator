"""Structured DOCX extraction preserving paths to paragraphs, runs, tables, and headers."""

from __future__ import annotations

import io
import logging
from collections.abc import Callable

from ..config.settings import Settings
from ..exceptions import CorruptedDocumentError, DependencyUnavailableError
from ..schemas import (
    BlockType,
    BoundingBox,
    ContentKind,
    DocumentModel,
    FileFormat,
    FontMetadata,
    PageModel,
    ProcessingOptions,
    Region,
    RegionType,
    TextBlock,
)
from ..utils.text_utils import normalize_text
from ..utils.validation import ValidatedFile
from .image_processor import ImagePreprocessor, load_image
from .page_ocr import PageOCRPipeline

LOGGER = logging.getLogger(__name__)


def _font_from_paragraph(paragraph: object) -> FontMetadata:
    run = next((item for item in paragraph.runs if item.text.strip()), paragraph.runs[0] if paragraph.runs else None)
    if run is None:
        return FontMetadata()
    size = float(run.font.size.pt) if run.font.size else 11.0
    color = "#000000"
    try:
        if run.font.color and run.font.color.rgb:
            color = f"#{run.font.color.rgb}"
    except Exception:
        LOGGER.debug("DOCX run color could not be read", exc_info=True)
    return FontMetadata(
        family=run.font.name or "Calibri",
        size=size,
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        color=color,
    )


def _paragraph_type(paragraph: object, region: str) -> BlockType:
    if region == "header":
        return BlockType.HEADER
    if region == "footer":
        return BlockType.FOOTER
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "")).casefold()
    if "title" in style_name:
        return BlockType.TITLE
    if "heading" in style_name:
        return BlockType.HEADING
    return BlockType.PARAGRAPH


class DOCXProcessor:
    def __init__(
        self,
        settings: Settings,
        page_ocr_loader: Callable[[], PageOCRPipeline],
        preprocessor: ImagePreprocessor,
    ) -> None:
        self.settings = settings
        self.page_ocr_loader = page_ocr_loader
        self.preprocessor = preprocessor

    @staticmethod
    def _block_from_paragraph(
        paragraph: object,
        reference: str,
        index: int,
        region: str = "body",
    ) -> TextBlock | None:
        text = normalize_text(paragraph.text)
        if not text:
            return None
        top = 54.0 + index * 18.0
        return TextBlock(
            page_number=1,
            block_type=_paragraph_type(paragraph, region),
            source_bbox=BoundingBox(54.0, top, 558.0, top + 16.0),
            source_text=text,
            normalized_text=text,
            font=_font_from_paragraph(paragraph),
            alignment=str(getattr(paragraph.alignment, "name", "left") or "left").lower(),
            source_reference=reference,
            region_type=RegionType.PRINTED_TEXT,
            ocr_engine="python-docx_native",
            provenance=["Native DOCX text extracted with run/style references"],
            metadata={
                "origin": "native_docx",
                "style": getattr(getattr(paragraph, "style", None), "name", None),
                "runs": [run.text for run in paragraph.runs],
                "native_source_snapshot": text,
                "native_run_count": len(paragraph.runs),
                "paragraph_format": {
                    "left_indent": int(paragraph.paragraph_format.left_indent or 0),
                    "right_indent": int(paragraph.paragraph_format.right_indent or 0),
                    "first_line_indent": int(
                        paragraph.paragraph_format.first_line_indent or 0
                    ),
                    "space_before": int(paragraph.paragraph_format.space_before or 0),
                    "space_after": int(paragraph.paragraph_format.space_after or 0),
                    "line_spacing": str(paragraph.paragraph_format.line_spacing or ""),
                    "keep_together": paragraph.paragraph_format.keep_together,
                    "keep_with_next": paragraph.paragraph_format.keep_with_next,
                    "page_break_before": paragraph.paragraph_format.page_break_before,
                },
            },
        )

    def _extract_native(self, document: object) -> list[TextBlock]:
        blocks: list[TextBlock] = []
        order = 0
        for paragraph_index, paragraph in enumerate(document.paragraphs):
            block = self._block_from_paragraph(paragraph, f"body:p:{paragraph_index}", order)
            if block:
                blocks.append(block)
                order += 1
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        reference = f"table:{table_index}:{row_index}:{cell_index}:p:{paragraph_index}"
                        block = self._block_from_paragraph(paragraph, reference, order)
                        if block:
                            block.block_type = BlockType.TABLE_CELL
                            block.region_type = RegionType.TABLE_FORM
                            block.metadata["table_coordinates"] = [table_index, row_index, cell_index]
                            blocks.append(block)
                            order += 1
        for section_index, section in enumerate(document.sections):
            for region_name, region in (("header", section.header), ("footer", section.footer)):
                for paragraph_index, paragraph in enumerate(region.paragraphs):
                    reference = f"section:{section_index}:{region_name}:p:{paragraph_index}"
                    block = self._block_from_paragraph(paragraph, reference, order, region_name)
                    if block:
                        blocks.append(block)
                        order += 1
        return blocks

    def _extract_images(
        self, document: object, options: ProcessingOptions, start_order: int
    ) -> tuple[list[TextBlock], list[str]]:
        blocks: list[TextBlock] = []
        warnings: list[str] = []
        seen_hashes: set[int] = set()
        order = start_order
        for relationship_id, relationship in document.part.rels.items():
            target = getattr(relationship, "target_part", None)
            content_type = str(getattr(target, "content_type", ""))
            if not content_type.startswith("image/"):
                continue
            blob = bytes(getattr(target, "blob", b""))
            if not blob or hash(blob) in seen_hashes:
                continue
            seen_hashes.add(hash(blob))
            try:
                image = load_image(blob)
                processed = (
                    self.preprocessor.preprocess(
                        image,
                        profile=options.preprocessing_profile,
                        allow_geometry=False,
                        upscale_factor=options.ocr_upscale_factor,
                    )
                    if options.enable_preprocessing
                    else None
                )
                ocr_image = processed.image if processed else image
                display_image = processed.display_image if processed else image
                ocr = self.page_ocr_loader().process(
                    ocr_image,
                    display_image,
                    page_number=1,
                    page_width=504.0,
                    page_height=max(72.0, image.height * 504.0 / max(1, image.width)),
                    options=options,
                    ocr_variants=(
                        processed.candidate_images
                        if processed
                        else {"original": image.convert("L")}
                    ),
                    quality_metrics=(
                        processed.quality.to_dict()
                        if processed and processed.quality
                        else {}
                    ),
                )
                warnings.extend(ocr.warnings)
                for block in ocr.blocks:
                    offset = 54.0 + order * 18.0
                    height = max(16.0, block.source_bbox.height)
                    block.source_bbox = BoundingBox(54.0, offset, 558.0, offset + height)
                    block.source_reference = f"embedded:{relationship_id}"
                    block.metadata["append_after_embedded_image"] = True
                    block.metadata["preserve_region_as_image"] = True
                    block.metadata["primary_output_replaceable"] = False
                    blocks.append(block)
                    order += 1
            except Exception as exc:
                warnings.append(f"Embedded image {relationship_id} could not be OCR processed.")
                LOGGER.warning("DOCX embedded-image OCR failed: %s", type(exc).__name__, exc_info=True)
        return blocks, warnings

    def process(
        self,
        validated: ValidatedFile,
        options: ProcessingOptions,
        stage_callback: Callable[[str, float], None] | None = None,
    ) -> DocumentModel:
        try:
            from docx import Document
        except ImportError as exc:
            raise DependencyUnavailableError("python-docx is required for Word documents") from exc
        try:
            document = Document(io.BytesIO(validated.data))
        except Exception as exc:
            raise CorruptedDocumentError("python-docx could not open the document") from exc
        native = self._extract_native(document)
        image_blocks: list[TextBlock] = []
        warnings: list[str] = []
        if document.inline_shapes:
            if stage_callback:
                stage_callback("ocr", 0.5)
            image_blocks, warnings = self._extract_images(document, options, len(native))
        if not native and not image_blocks and not document.inline_shapes:
            raise CorruptedDocumentError("DOCX contains no text or images")
        content_kind = (
            ContentKind.MIXED
            if native and image_blocks
            else ContentKind.SCANNED
            if image_blocks or document.inline_shapes and not native
            else ContentKind.NATIVE
        )
        blocks = native + image_blocks
        regions: list[Region] = []
        for order, block in enumerate(blocks):
            region = Region(
                page_number=1,
                bbox=block.source_bbox,
                region_type=block.region_type,
                classification_confidence=1.0 if not block.is_ocr else 0.7,
                reading_order=order,
                block_ids=[block.block_id],
                preserve_as_image=bool(block.metadata.get("append_after_embedded_image")),
                overlaps_critical_graphic=block.region_type
                in {RegionType.SIGNATURE, RegionType.STAMP_SEAL},
                metadata={"source_reference": block.source_reference},
            )
            block.metadata["region_id"] = region.region_id
            block.metadata["region_reading_order"] = order
            regions.append(region)
        page_height = max(792.0, 90.0 + len(blocks) * 18.0)
        return DocumentModel(
            document_id=validated.sha256[:20],
            source_filename=validated.filename,
            file_format=FileFormat.DOCX,
            content_kind=content_kind,
            pages=[
                PageModel(
                    1,
                    612.0,
                    page_height,
                    blocks,
                    regions=regions,
                    content_kind=content_kind,
                )
            ],
            source_bytes=validated.data,
            mime_type=validated.mime_type,
            warnings=list(dict.fromkeys(warnings)),
            metadata={"section_count": len(document.sections), "inline_shape_count": len(document.inline_shapes)},
        )
