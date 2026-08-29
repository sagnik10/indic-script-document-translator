"""Model-free tests for the Kaggle handwritten-page inference application."""

from __future__ import annotations

import contextlib
import re
import sys
import types
from io import BytesIO

import cv2
import numpy as np
from PIL import Image
import pytest

import pretrained_page_ocr.kaggle_app as page_app
from pretrained_page_ocr.kaggle_app import (
    EnhancedPage,
    LineResult,
    ModelRegistry,
    RecognitionCandidate,
    TextRegion,
    _find_tessdata_directory,
    _find_tesseract_executable,
    _finish_pretrained_model_load,
    _opencv_detect,
    _split_context_translation,
    _square_pad,
    _word_document_line,
    build_translated_docx,
    calculate_text_quality,
    build_translation_units,
    choose_best_recognition,
    detect_script,
    enhance_page,
    merge_text_regions,
    normalize_translation_source,
    protect_translation_source,
    restore_translation_source,
    save_results,
    script_purity,
    script_statistics,
    sort_reading_order,
    translation_quality,
    translate_text,
)


def test_tesseract_paths_from_explicit_overrides(monkeypatch, tmp_path) -> None:
    executable_path = tmp_path / "tesseract.exe"
    executable_path.write_bytes(b"test executable")
    tessdata_path = tmp_path / "tessdata"
    tessdata_path.mkdir()
    for language in ("eng", "hin", "pan", "osd"):
        (tessdata_path / f"{language}.traineddata").write_bytes(b"test data")
    monkeypatch.setenv("TESSERACT_CMD", str(executable_path))
    monkeypatch.setenv("TESSDATA_PREFIX", str(tessdata_path))

    executable = _find_tesseract_executable()
    tessdata = _find_tessdata_directory(("eng", "hin", "pan", "osd"))

    assert executable == executable_path.resolve()
    assert tessdata == tessdata_path.resolve()
    assert (tessdata / "hin.traineddata").is_file()
    assert (tessdata / "pan.traineddata").is_file()


def test_hindi_loader_builds_legacy_checkpoint_processor_explicitly(monkeypatch) -> None:
    calls: dict[str, object] = {}

    class FakeImageProcessor:
        @classmethod
        def from_pretrained(cls, model_id: str, **_kwargs):
            calls["image_model_id"] = model_id
            return "image-processor"

    class FakeTokenizer:
        @classmethod
        def from_pretrained(cls, model_id: str, **_kwargs):
            calls["tokenizer_model_id"] = model_id
            return "tokenizer"

    class FakeProcessor:
        def __init__(self, *, image_processor, tokenizer):
            calls["processor"] = (image_processor, tokenizer)

    class FakeModel:
        @classmethod
        def from_pretrained(cls, model_id: str, **kwargs):
            calls["model_id"] = model_id
            calls["model_load_options"] = kwargs
            return cls()

        def to(self, device):
            calls["device"] = str(device)
            return self

        def eval(self):
            calls["eval"] = True

    registry = ModelRegistry()
    fake_transformers = types.ModuleType("transformers")
    fake_transformers.ViTImageProcessorPil = FakeImageProcessor
    fake_transformers.RobertaTokenizer = FakeTokenizer
    fake_transformers.TrOCRProcessor = FakeProcessor
    fake_transformers.VisionEncoderDecoderModel = FakeModel
    monkeypatch.setitem(sys.modules, "transformers", fake_transformers)
    processor, model = registry.load_hindi()

    assert isinstance(processor, FakeProcessor)
    assert isinstance(model, FakeModel)
    assert calls["processor"] == ("image-processor", "tokenizer")
    assert calls["eval"] is True
    assert calls["model_load_options"]["low_cpu_mem_usage"] is False


def test_meta_checkpoint_is_rejected_without_random_materialization() -> None:
    import torch

    model = torch.nn.Linear(4, 3, device="meta")

    with pytest.raises(RuntimeError, match="unmaterialized meta tensors"):
        _finish_pretrained_model_load(model, torch.device("cpu"))

    assert model.weight.is_meta is True


def test_gurmukhi_loader_forces_float32_on_cpu(monkeypatch) -> None:
    calls: dict[str, object] = {}

    class FakeTokenizer:
        padding_side = "right"
        pad_token_id = None
        eos_token_id = 2

    class FakeProcessor:
        tokenizer = FakeTokenizer()

        @classmethod
        def from_pretrained(cls, model_id: str, **_kwargs):
            calls["processor_model_id"] = model_id
            return cls()

    class FakeModel:
        generation_config = types.SimpleNamespace(
            eos_token_id=None,
            pad_token_id=None,
        )

        @classmethod
        def from_pretrained(cls, model_id: str, **kwargs):
            calls["model_id"] = model_id
            calls["dtype"] = kwargs.get("dtype")
            return cls()

        def to(self, device):
            calls["device"] = str(device)
            return self

        def eval(self):
            calls["eval"] = True

    registry = ModelRegistry()
    registry.device = registry.torch.device("cpu")
    registry.dtype = registry.torch.float32
    fake_transformers = types.ModuleType("transformers")
    fake_transformers.AutoProcessor = FakeProcessor
    fake_transformers.AutoModelForMultimodalLM = FakeModel
    monkeypatch.setitem(sys.modules, "transformers", fake_transformers)

    processor, model = registry.load_gurmukhi()

    assert isinstance(processor, FakeProcessor)
    assert isinstance(model, FakeModel)
    assert calls["dtype"] is registry.torch.float32
    assert calls["device"] == "cpu"
    assert calls["eval"] is True
    assert model.generation_config.max_length is None


def _candidate(
    script: str,
    text: str,
    confidence: float = 0.85,
    *,
    model_id: str = "unit-test-model",
) -> RecognitionCandidate:
    return RecognitionCandidate(
        expected_script=script,  # type: ignore[arg-type]
        text=text,
        confidence=confidence,
        model_id=model_id,
        provider_kind="unit_test",
    )


def test_script_statistics_detect_gurmukhi_and_ignore_digits() -> None:
    text = "ਪੰਜਾਬੀ ਭਾਸ਼ਾ 2026"

    stats = script_statistics(text)

    # Combining vowel/sign code points are intentionally excluded because the
    # application defines its denominator as ``str.isalpha()`` characters.
    assert stats["gurmukhi"] == 5
    assert stats["devanagari"] == 0
    assert stats["letters"] == 5
    assert stats["gurmukhi_ratio"] == 1.0
    assert detect_script(text) == "gurmukhi"
    assert script_purity(text, "gurmukhi") == 1.0


def test_script_statistics_detect_devanagari_and_mixed_text() -> None:
    hindi = "हिंदी भाषा"
    mixed = "हिंदी ਪੰਜਾਬੀ"

    assert detect_script(hindi) == "devanagari"
    assert script_statistics(hindi)["devanagari_ratio"] == 1.0
    assert detect_script(mixed) == "mixed"
    assert 0.0 < script_purity(mixed, "mixed") <= 1.0
    assert detect_script("2026 | ?") == "unknown"


def test_candidate_choice_prefers_script_consistent_source_text() -> None:
    candidates = [
        _candidate("gurmukhi", "ਪੰਜਾਬੀ ਭਾਸ਼ਾ ਦਾ ਦਸਤਾਵੇਜ਼", 0.81),
        _candidate("devanagari", "? || a", 0.98),
    ]

    winner, accepted, reasons = choose_best_recognition(candidates)

    assert winner.expected_script == "gurmukhi"
    assert winner.detected_script == "gurmukhi"
    assert winner.text == "ਪੰਜਾਬੀ ਭਾਸ਼ਾ ਦਾ ਦਸਤਾਵੇਜ਼"
    assert accepted is True
    assert reasons == []


def test_candidate_choice_marks_wrong_script_and_garbage_for_review() -> None:
    winner, accepted, reasons = choose_best_recognition(
        [_candidate("gurmukhi", "I | ??", 0.95)]
    )

    assert winner.detected_script == "unknown"
    assert accepted is False
    assert "too few source-script letters" in reasons
    assert "low Unicode script purity" in reasons
    assert "model/script disagreement" in reasons


def test_missing_confidence_and_unusually_short_text_require_review() -> None:
    no_confidence = _candidate("gurmukhi", "ਪੰਜਾਬੀ ਲਿਖਤ", 0.8)
    no_confidence.confidence = None
    _winner, accepted, reasons = choose_best_recognition([no_confidence])
    assert accepted is False
    assert "OCR confidence unavailable" in reasons

    _winner, accepted, reasons = choose_best_recognition(
        [_candidate("devanagari", "यह", 0.92)]
    )
    assert accepted is False
    assert "unusually short OCR output" in reasons


def test_square_padding_preserves_word_aspect_ratio() -> None:
    word = np.full((30, 120, 3), 255, dtype=np.uint8)
    word[8:22, 20:100] = 0
    padded = _square_pad(Image.fromarray(word))

    assert padded.width == padded.height
    assert padded.width == 136


def test_text_quality_rewards_source_script_and_rejects_symbols() -> None:
    gurmukhi = calculate_text_quality(
        "ਇਹ ਇੱਕ ਪੜ੍ਹਨ ਯੋਗ ਪੰਜਾਬੀ ਵਾਕ ਹੈ", "gurmukhi"
    )
    devanagari = calculate_text_quality(
        "यह एक पठनीय हिंदी वाक्य है", "devanagari"
    )
    garbage = calculate_text_quality("||| ??? £", "gurmukhi")
    wrong_script = calculate_text_quality("This is Latin", "gurmukhi")

    assert gurmukhi >= 0.75
    assert devanagari >= 0.75
    assert garbage < 0.20
    assert wrong_script < gurmukhi


def test_merge_text_regions_joins_fragments_on_same_baseline() -> None:
    regions = [
        TextRegion((20, 40, 100, 62), 0.80, "test"),
        TextRegion((112, 41, 220, 64), 0.90, "test"),
        TextRegion((25, 105, 210, 130), 0.75, "test"),
    ]

    merged = merge_text_regions(regions)

    assert len(merged) == 2
    assert merged[0].bbox == (20, 40, 220, 64)
    assert abs((merged[0].detection_confidence or 0.0) - 0.85) < 1e-9
    assert merged[1].bbox == (25, 105, 210, 130)


def test_reading_order_is_top_to_bottom_then_left_to_right() -> None:
    regions = [
        TextRegion((210, 25, 310, 48), None, "test"),
        TextRegion((25, 100, 200, 124), None, "test"),
        TextRegion((20, 27, 180, 50), None, "test"),
        TextRegion((28, 165, 230, 191), None, "test"),
    ]

    ordered = sort_reading_order(regions, page_width=400)

    assert [region.bbox for region in ordered] == [
        (20, 27, 180, 50),
        (210, 25, 310, 48),
        (25, 100, 200, 124),
        (28, 165, 230, 191),
    ]


def _synthetic_page() -> EnhancedPage:
    height, width = 420, 640
    threshold = np.full((height, width), 255, dtype=np.uint8)

    # Three uneven, handwriting-like ink lines. Short strokes avoid looking like
    # table rules to the fallback detector.
    for row, y in enumerate((80, 185, 292)):
        x = 65 + row * 7
        for index in range(12):
            stroke_width = 18 + (index % 4) * 5
            top = y + ((index % 3) - 1) * 3
            cv2.ellipse(
                threshold,
                (x + stroke_width // 2, top),
                (stroke_width // 2, 8 + index % 3),
                0,
                0,
                300,
                0,
                3,
            )
            x += stroke_width + 7

    rgb = cv2.cvtColor(threshold, cv2.COLOR_GRAY2RGB)
    return EnhancedPage(
        corrected_rgb=rgb,
        enhanced_rgb=rgb.copy(),
        enhanced_gray=threshold.copy(),
        threshold=threshold,
        operations=["synthetic_test_page"],
    )


def test_opencv_fallback_detects_logical_lines_on_synthetic_page() -> None:
    page = _synthetic_page()

    regions = _opencv_detect(page)

    assert 2 <= len(regions) <= 4
    assert [region.bbox[1] for region in regions] == sorted(
        region.bbox[1] for region in regions
    )
    assert all(region.detector == "opencv_projection_fallback" for region in regions)
    assert all(region.bbox[0] > 0 and region.bbox[2] < 640 for region in regions)


def test_enhancement_preserves_faint_dark_handwriting_contrast() -> None:
    gray = np.full((220, 480), 225, dtype=np.uint8)
    cv2.putText(
        gray,
        "faint line",
        (45, 120),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.2,
        118,
        2,
        cv2.LINE_AA,
    )
    rgb = cv2.cvtColor(gray, cv2.COLOR_GRAY2RGB)

    enhanced = enhance_page(rgb)

    ink_mask = gray < 180
    paper_mask = gray > 215
    contrast = float(enhanced.enhanced_gray[paper_mask].mean()) - float(
        enhanced.enhanced_gray[ink_mask].mean()
    )
    assert contrast >= 35.0


def _line_result(
    line: int,
    *,
    text: str = "",
    english: str = "",
    accepted: bool = False,
    translation_status: str = "not_attempted",
) -> LineResult:
    return LineResult(
        line=line,
        bbox=(40, 30 + line * 40, 520, 55 + line * 40),
        crop_file=f"line_crops/line_{line:03d}.jpg",
        detection_confidence=0.9,
        detector="unit_test",
        script="gurmukhi" if text else "unknown",
        text=text,
        script_purity=1.0 if text else 0.0,
        confidence=0.9 if text else None,
        text_quality=0.9 if text else 0.0,
        recognition_model="unit-test-model",
        provider_kind="unit_test",
        accepted=accepted,
        review_required=not (accepted and translation_status == "translated"),
        review_reasons=[] if accepted else ["test review"],
        english=english,
        translation_status=translation_status,
    )


def test_translated_word_document_contains_every_detected_line_in_order() -> None:
    from docx import Document

    lines = [
        _line_result(
            1,
            text="ਪੰਜਾਬੀ ਸਰੋਤ",
            english="English translation",
            accepted=True,
            translation_status="translated",
        ),
        _line_result(2, text="ਅਨਿਸ਼ਚਿਤ ਸਰੋਤ"),
        _line_result(3),
    ]

    content = build_translated_docx(_synthetic_page(), lines)
    document = Document(BytesIO(content))

    assert content[:2] == b"PK"
    assert [paragraph.text for paragraph in document.paragraphs] == [
        _word_document_line(line) for line in lines
    ]
    assert document.paragraphs[0].text == "English translation"
    assert "ਅਨਿਸ਼ਚਿਤ ਸਰੋਤ" in document.paragraphs[1].text
    assert "no reliable transcription" in document.paragraphs[2].text


def test_save_results_writes_primary_word_artifact(tmp_path) -> None:
    lines = [
        _line_result(
            1,
            text="ਪੰਜਾਬੀ ਸਰੋਤ",
            english="Translated line",
            accepted=True,
            translation_status="translated",
        )
    ]

    payload = save_results(
        input_image=tmp_path / "source.jpg",
        page=_synthetic_page(),
        lines=lines,
        models=ModelRegistry(),
        timings={"total": 0.1},
        output_directory=tmp_path / "output",
    )

    output = tmp_path / "output" / "translated_en.docx"
    assert payload["primary_output"] == "translated_en.docx"
    assert output.read_bytes()[:2] == b"PK"


def test_multi_page_word_validation_ignores_structural_page_break_paragraphs() -> None:
    from docx import Document

    first = _line_result(
        1,
        text="ਪਹਿਲਾ ਪੰਨਾ",
        english="First page",
        accepted=True,
        translation_status="translated",
    )
    second = _line_result(
        2,
        text="ਦੂਜਾ ਪੰਨਾ",
        english="Second page",
        accepted=True,
        translation_status="translated",
    )
    second.page_number = 2

    content = build_translated_docx(_synthetic_page(), [first, second])
    reopened = Document(BytesIO(content))

    assert [paragraph.text for paragraph in reopened.paragraphs if paragraph.text] == [
        "First page",
        "Second page",
    ]


def test_fast_cpu_mode_uses_batched_gurmukhi_and_quality_translation(
    monkeypatch, tmp_path
) -> None:
    calls: dict[str, object] = {}

    class FakeDevice:
        type = "cpu"

    class FakeModels:
        device = FakeDevice()
        gpu_name = "CPU"
        status = {
            "detector": {"loaded": False},
            "hindi": {"id": page_app.HINDI_MODEL_ID, "loaded": False},
            "gurmukhi": {"id": page_app.GURMUKHI_MODEL_ID, "loaded": False},
            "translation": {"id": page_app.TRANSLATION_MODEL_ID, "loaded": False},
        }

        def release_recognizers(self) -> None:
            calls["released"] = True

    source_path = tmp_path / "source.png"
    Image.new("RGB", (240, 160), "white").save(source_path)
    monkeypatch.setattr(page_app, "load_models", lambda: FakeModels())
    monkeypatch.setattr(
        page_app,
        "detect_text_regions",
        lambda page, models: [TextRegion((20, 40, 220, 75), 0.9, "unit_test")],
    )

    def fake_gurmukhi(crops, models, progress_callback=None, *, fast=False):
        calls["fast"] = fast
        return [_candidate("gurmukhi", "ਪੰਜਾਬੀ ਭਾਸ਼ਾ ਦੀ ਲਿਖਤ", 0.95)]

    def fake_translate(lines, models, progress_callback=None, num_beams=4):
        calls["num_beams"] = num_beams
        for line in lines:
            line.english = "Punjabi language text"
            line.translation_status = "translated"
            line.review_required = False
        return 0.01

    monkeypatch.setattr(page_app, "recognize_gurmukhi", fake_gurmukhi)
    monkeypatch.setattr(page_app, "translate_text", fake_translate)

    payload = page_app.process_page(
        source_path,
        tmp_path / "output",
        processing_mode="fast_cpu",
    )

    assert calls == {"fast": True, "released": True, "num_beams": 4}
    assert payload["lines"][0]["english"] == "Punjabi language text"
    assert (tmp_path / "output" / "translated_en.docx").is_file()


def test_translation_source_normalization_only_repairs_combining_spacing() -> None:
    assert normalize_translation_source("ਪੰਜਾਬ ੀ ਭਾਸ਼ਾ") == "ਪੰਜਾਬੀ ਭਾਸ਼ਾ"
    assert normalize_translation_source("हिंद ी भाषा") == "हिंदी भाषा"
    assert normalize_translation_source("ਪੰਜਾਬੀ ਭਾਸ਼ਾ") == "ਪੰਜਾਬੀ ਭਾਸ਼ਾ"


def test_translation_protection_restores_numbers_ids_and_existing_english() -> None:
    source = "ਮਿਤੀ 31.03.2016 ਨੂੰ Civil Hospital Ref No. PB-204 ਭੇਜਿਆ"
    protected = protect_translation_source(source)

    assert "31.03.2016" not in protected.text
    assert "Civil Hospital" not in protected.text
    assert "PB-204" not in protected.text

    translated = protected.text.replace("ਮਿਤੀ", "Date").replace("ਨੂੰ", "on")
    restored, missing = restore_translation_source(translated, protected)

    assert not missing
    assert "31.03.2016" in restored
    assert "Civil Hospital" in restored
    assert "PB-204" in restored


def test_context_units_group_only_adjacent_validated_scan_lines() -> None:
    first = _line_result(1, text="ਪੰਜਾਬੀ ਸਰੋਤ ਦੀ ਪਹਿਲੀ ਲਾਈਨ", accepted=True)
    second = _line_result(2, text="ਅਤੇ ਇਹ ਅਗਲੀ ਲਾਈਨ ਹੈ", accepted=True)
    rejected = _line_result(3, text="ਅਸਪਸ਼ਟ")
    fourth = _line_result(4, text="ਇਹ ਵੱਖਰਾ ਵਾਕ ਹੈ।", accepted=True)

    units = build_translation_units([first, second, rejected, fourth])

    assert [len(unit.lines) for unit in units] == [2, 1]
    assert units[0].source == f"{first.text} {second.text}"


def test_context_translation_split_preserves_exact_protected_phrase() -> None:
    first = _line_result(1, text="ਪਹਿਲੀ ਪੰਜਾਬੀ ਲਾਈਨ", accepted=True)
    second = _line_result(2, text="ਦੂਜੀ ਪੰਜਾਬੀ ਲਾਈਨ", accepted=True)

    values = _split_context_translation(
        "The case was sent to Civil Hospital on 31.03.2016 for review",
        [first, second],
        ["Civil Hospital", "31.03.2016"],
    )

    assert values is not None
    assert " ".join(values) == (
        "The case was sent to Civil Hospital on 31.03.2016 for review"
    )
    assert any("Civil Hospital" in value for value in values)


def test_translation_quality_rejects_non_english_or_repetitive_output() -> None:
    source = "ਇਹ ਇੱਕ ਪੜ੍ਹਨ ਯੋਗ ਪੰਜਾਬੀ ਵਾਕ ਹੈ"

    good_score, good_reasons = translation_quality(
        source,
        "This is a readable Punjabi sentence.",
    )
    bad_score, bad_reasons = translation_quality(
        source,
        "ਹਾਂ ਹਾਂ ਹਾਂ ਹਾਂ ਹਾਂ ਹਾਂ ਹਾਂ ਹਾਂ",
    )

    assert good_score >= 0.55
    assert not good_reasons
    assert bad_score < good_score
    assert bad_reasons


def test_translation_uses_quality_beams_and_restores_exact_date(monkeypatch) -> None:
    import torch

    calls: dict[str, object] = {}

    class FakeTokenizer:
        src_lang = ""
        source: list[str] = []

        def convert_tokens_to_ids(self, token: str) -> int:
            assert token == "eng_Latn"
            return 7

        def __call__(self, texts, **_kwargs):
            self.source = list(texts)
            return {"input_ids": torch.tensor([[1, 2, 3]])}

        def batch_decode(self, _sequences, skip_special_tokens=True):
            assert skip_special_tokens is True
            placeholder = re.search(r"ZXQPROTECTED[A-Z]+QXZ", self.source[0])
            assert placeholder is not None
            return [f"The document dated {placeholder.group(0)} was received"]

    class FakeModel:
        def generate(self, **kwargs):
            calls.update(kwargs)
            return types.SimpleNamespace(
                sequences=torch.tensor([[1, 2, 3]]),
                scores=(),
                beam_indices=None,
            )

    class FakeModels:
        def __init__(self):
            self.torch = torch
            self.device = torch.device("cpu")
            self.status = {"translation": {"id": page_app.TRANSLATION_MODEL_ID}}

        def load_translation(self):
            return FakeTokenizer(), FakeModel()

        def release_translation(self):
            calls["released"] = True

        def autocast(self):
            return contextlib.nullcontext()

    monkeypatch.setattr(page_app, "_generation_confidences", lambda *_args: [0.91])
    line = _line_result(
        1,
        text="ਦਸਤਾਵੇਜ਼ ਮਿਤੀ 31.03.2016 ਨੂੰ ਮਿਲਿਆ",
        accepted=True,
    )

    translate_text([line], FakeModels(), num_beams=4)

    assert line.translation_status == "translated"
    assert "31.03.2016" in line.english
    assert line.translation_confidence == 0.91
    assert line.translation_quality >= 0.55
    assert calls["num_beams"] == 4
    assert calls["no_repeat_ngram_size"] == 3
    assert calls["released"] is True
