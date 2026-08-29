from __future__ import annotations

import io

from docx import Document
from docx.shared import Inches
from PIL import Image

from translator_app.config.settings import Settings
from translator_app.core.docx_processor import DOCXProcessor
from translator_app.core.image_processor import ImagePreprocessor
from translator_app.reconstruction.docx_reconstructor import reconstruct_docx
from translator_app.schemas import (
    LayoutStatus,
    ProcessingOptions,
    ReconstructionType,
    RegionType,
    TranslationStatus,
)
from translator_app.utils.validation import validate_upload


def test_docx_processor_and_reconstructor_preserve_runs_and_table() -> None:
    source = Document()
    paragraph = source.add_paragraph()
    paragraph.add_run("मूल ")
    bold_run = paragraph.add_run("पाठ")
    bold_run.bold = True
    table = source.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "नाम"
    buffer = io.BytesIO()
    source.save(buffer)
    settings = Settings(ocr_languages=["eng", "hin"])
    validated = validate_upload("regional.docx", buffer.getvalue(), 2_000_000)
    processor = DOCXProcessor(settings, lambda: object(), ImagePreprocessor())
    model = processor.process(validated, ProcessingOptions(ocr_languages=["eng", "hin"]))
    body_block = next(block for block in model.blocks if block.source_reference == "body:p:0")
    body_block.english_translation = "Original text"
    body_block.translation_status = TranslationStatus.TRANSLATED
    body_block.detected_language = "hi"
    body_block.source_validated = True
    cell_block = next(block for block in model.blocks if block.source_reference.startswith("table:"))
    cell_block.english_translation = "Name"
    cell_block.translation_status = TranslationStatus.TRANSLATED
    cell_block.detected_language = "hi"
    cell_block.source_validated = True
    output = reconstruct_docx(model)
    rebuilt = Document(io.BytesIO(output))
    assert "Original text" == " ".join(rebuilt.paragraphs[0].text.split())
    assert any(run.bold for run in rebuilt.paragraphs[0].runs)
    assert rebuilt.tables[0].cell(0, 0).text == "Name"
    assert body_block.metadata["replacement_applied"] is True
    assert cell_block.metadata["replacement_applied"] is True
    assert body_block.layout_status == LayoutStatus.FIT
    assert model.metadata["applied_translation_replacement_count"] == 2


def _mark_valid_hindi_translation(block: object, translation: str) -> None:
    block.english_translation = translation
    block.translation_status = TranslationStatus.TRANSLATED
    block.detected_language = "hi"
    block.source_validated = True


def test_primary_docx_preserves_media_geometry_and_does_not_append_diagnostics() -> None:
    image_bytes = io.BytesIO()
    Image.new("RGB", (24, 18), "navy").save(image_bytes, format="PNG")
    source = Document()
    section = source.sections[0]
    section.page_width = Inches(7.5)
    section.page_height = Inches(10.0)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.9)
    paragraph = source.add_paragraph()
    paragraph.style = source.styles["Title"]
    paragraph.add_run("मूल ")
    image_run = paragraph.add_run()
    image_run.add_picture(io.BytesIO(image_bytes.getvalue()), width=Inches(0.25))
    emphasized = paragraph.add_run("पाठ")
    emphasized.italic = True
    source.add_paragraph("अस्पष्ट पंक्ति")
    source.sections[0].header.paragraphs[0].text = "Hospital header"
    source.sections[0].footer.paragraphs[0].text = "Page footer"
    buffer = io.BytesIO()
    source.save(buffer)

    settings = Settings(ocr_languages=["eng", "hin"])
    validated = validate_upload("formatted.docx", buffer.getvalue(), 2_000_000)
    processor = DOCXProcessor(settings, lambda: object(), ImagePreprocessor())
    model = processor.process(validated, ProcessingOptions(ocr_languages=["eng", "hin"]))
    block = next(item for item in model.blocks if item.source_reference == "body:p:0")
    _mark_valid_hindi_translation(block, "Original translated text")
    model.metadata["reconstruction_mode"] = "overlay_translation"

    output = reconstruct_docx(model)
    rebuilt = Document(io.BytesIO(output))

    assert len(rebuilt.paragraphs) == len(source.paragraphs)
    assert len(rebuilt.sections) == len(source.sections)
    assert len(rebuilt.inline_shapes) == 1
    assert rebuilt.sections[0].page_width == source.sections[0].page_width
    assert rebuilt.sections[0].page_height == source.sections[0].page_height
    assert rebuilt.sections[0].left_margin == source.sections[0].left_margin
    assert rebuilt.sections[0].right_margin == source.sections[0].right_margin
    assert rebuilt.sections[0].header.paragraphs[0].text == "Hospital header"
    assert rebuilt.sections[0].footer.paragraphs[0].text == "Page footer"
    assert rebuilt.paragraphs[0].style.name == "Title"
    assert any(run.italic for run in rebuilt.paragraphs[0].runs)
    assert "Original translated text" == " ".join(rebuilt.paragraphs[0].text.split())
    full_text = "\n".join(paragraph.text for paragraph in rebuilt.paragraphs)
    assert "AI processing audit" not in full_text
    assert "English translation annotations" not in full_text
    assert "Translated text from embedded images" not in full_text
    assert model.metadata["docx_structure_validated"] is True
    assert model.metadata["docx_primary_output_has_appended_diagnostics"] is False


def test_docx_critical_and_unreadable_regions_remain_unchanged() -> None:
    source = Document()
    source.add_paragraph("मुहर")
    source.add_paragraph("अस्पष्ट लिखाई")
    source.add_paragraph("मूल पाठ")
    buffer = io.BytesIO()
    source.save(buffer)
    settings = Settings(ocr_languages=["eng", "hin"])
    validated = validate_upload("preserved.docx", buffer.getvalue(), 2_000_000)
    model = DOCXProcessor(settings, lambda: object(), ImagePreprocessor()).process(
        validated, ProcessingOptions(ocr_languages=["eng", "hin"])
    )
    stamp, unreadable, unvalidated = [
        next(item for item in model.blocks if item.source_reference == f"body:p:{index}")
        for index in range(3)
    ]
    _mark_valid_hindi_translation(stamp, "Seal")
    stamp.region_type = RegionType.STAMP_SEAL
    _mark_valid_hindi_translation(unreadable, "Unreadable writing")
    unreadable.reconstruction_type = ReconstructionType.UNREADABLE
    unvalidated.english_translation = "Original text"
    unvalidated.translation_status = TranslationStatus.TRANSLATED
    unvalidated.detected_language = "hi"

    rebuilt = Document(io.BytesIO(reconstruct_docx(model)))

    assert [paragraph.text for paragraph in rebuilt.paragraphs] == [
        "मुहर",
        "अस्पष्ट लिखाई",
        "मूल पाठ",
    ]
    assert stamp.metadata["primary_output_replacement_reason"] == "critical_graphic_region"
    assert unreadable.metadata["primary_output_replacement_reason"] == "unreadable_source"
    assert unvalidated.metadata["primary_output_replacement_reason"] == "source_not_validated"
    assert stamp.metadata["replacement_applied"] is False
    assert unreadable.metadata["replacement_applied"] is False
    assert unvalidated.metadata["replacement_applied"] is False
    assert stamp.layout_status == LayoutStatus.SKIPPED
    assert unreadable.layout_status == LayoutStatus.SKIPPED
    assert unvalidated.layout_status == LayoutStatus.SKIPPED
    assert model.metadata["applied_translation_replacement_count"] == 0
