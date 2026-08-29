"""Fast, structure-preserving DOCX transcription and Indic-to-English translation.

Native Word text is read directly; no OCR model is loaded for it.  When a Word
file contains no native Gurmukhi/Devanagari text but does contain embedded page
images, those images receive one local Tesseract ``pan+hin+eng`` pass.  Only
validated Gurmukhi/Devanagari source spans are sent to the existing NLLB
translation helper.  Existing English, numbers, identifiers, and punctuation
are copied byte-for-text into reconstructed native paragraphs.
"""

from __future__ import annotations

import hashlib
import json
import re
import time
from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, Iterator, Literal, Sequence

import cv2
import numpy as np
from PIL import Image, ImageOps, UnidentifiedImageError

from pretrained_page_ocr.kaggle_app import (
    DEVANAGARI_END,
    DEVANAGARI_START,
    GURMUKHI_END,
    GURMUKHI_START,
    LineResult,
    ModelRegistry,
    ProgressCallback,
    RecognitionCandidate,
    _candidate_for_audit,
    _prepare_tesseract,
    _report_progress,
    choose_best_recognition,
    detect_script,
    enrich_candidate,
    normalize_text,
    recognize_gurmukhi,
    recognize_hindi,
    script_statistics,
    translate_text,
)


DocxProcessingMode = Literal["fast_document", "line_accurate"]


@dataclass(slots=True)
class _TextPart:
    """One preserved or source-language part of a native Word paragraph."""

    text: str
    script: Literal["gurmukhi", "devanagari"] | None = None
    line_result: LineResult | None = None


@dataclass(slots=True)
class _NativeUnit:
    """A mutable Word paragraph and its stable structural reference."""

    order: int
    reference: str
    paragraph: Any
    text: str
    parts: list[_TextPart]


def _character_script(character: str) -> str | None:
    codepoint = ord(character)
    if GURMUKHI_START <= codepoint <= GURMUKHI_END:
        return "gurmukhi"
    # U+0964/U+0965 are common Indic danda punctuation, not sufficient proof
    # of Devanagari language on their own.
    if DEVANAGARI_START <= codepoint <= DEVANAGARI_END and codepoint not in {
        0x0964,
        0x0965,
    }:
        return "devanagari"
    return None


def split_translatable_parts(text: str) -> list[_TextPart]:
    """Split text into source-script spans and exact pass-through spans.

    Numbers, Latin text, email/URL fragments, and punctuation are deliberately
    kept outside translation.  Whitespace bridges adjacent words in the same
    script so NLLB still receives useful phrase context instead of isolated
    characters or words.
    """

    if not text:
        return []
    runs: list[_TextPart] = []
    start = 0
    current_script = _character_script(text[0])
    for index, character in enumerate(text[1:], start=1):
        script = _character_script(character)
        if script == current_script:
            continue
        runs.append(_TextPart(text[start:index], current_script))
        start = index
        current_script = script
    runs.append(_TextPart(text[start:], current_script))

    # Neutral punctuation between same-script runs belongs to that phrase.
    # Numbers, Latin text, and identifiers remain protected and untouched.
    for index in range(1, len(runs) - 1):
        run = runs[index]
        if (
            run.script is None
            and re.fullmatch(r"[\s.,;:!?()'\"\-।॥]+", run.text)
            and runs[index - 1].script is not None
            and runs[index - 1].script == runs[index + 1].script
        ):
            run.script = runs[index - 1].script

    merged: list[_TextPart] = []
    for run in runs:
        if merged and merged[-1].script == run.script:
            merged[-1].text += run.text
        else:
            merged.append(run)
    return merged


def _iter_table_paragraphs(
    table: Any,
    table_reference: str,
    seen_cells: set[int],
) -> Iterator[tuple[str, Any]]:
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            identity = id(cell._tc)
            if identity in seen_cells:
                continue
            seen_cells.add(identity)
            cell_reference = f"{table_reference}:r:{row_index}:c:{cell_index}"
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield f"{cell_reference}:p:{paragraph_index}", paragraph
            for nested_index, nested in enumerate(cell.tables):
                yield from _iter_table_paragraphs(
                    nested,
                    f"{cell_reference}:table:{nested_index}",
                    seen_cells,
                )


def _iter_document_paragraphs(document: Any) -> Iterator[tuple[str, Any]]:
    """Yield body, table, header, and footer paragraphs without cell duplicates."""

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        yield f"body:p:{paragraph_index}", paragraph
    seen_cells: set[int] = set()
    for table_index, table in enumerate(document.tables):
        yield from _iter_table_paragraphs(table, f"body:table:{table_index}", seen_cells)
    seen_regions: set[int] = set()
    for section_index, section in enumerate(document.sections):
        for region_name, region in (
            ("header", section.header),
            ("footer", section.footer),
        ):
            region_identity = id(region.part)
            if region_identity in seen_regions:
                continue
            seen_regions.add(region_identity)
            for paragraph_index, paragraph in enumerate(region.paragraphs):
                yield (
                    f"section:{section_index}:{region_name}:p:{paragraph_index}",
                    paragraph,
                )
            for table_index, table in enumerate(region.tables):
                yield from _iter_table_paragraphs(
                    table,
                    f"section:{section_index}:{region_name}:table:{table_index}",
                    seen_cells,
                )


def _extract_native_units(document: Any) -> list[_NativeUnit]:
    units: list[_NativeUnit] = []
    for reference, paragraph in _iter_document_paragraphs(document):
        text = str(paragraph.text or "")
        if not text.strip():
            continue
        units.append(
            _NativeUnit(
                order=len(units),
                reference=reference,
                paragraph=paragraph,
                text=text,
                parts=split_translatable_parts(text),
            )
        )
    return units


def _native_line_result(
    *,
    line_number: int,
    reference: str,
    text: str,
    script: Literal["gurmukhi", "devanagari"],
) -> LineResult:
    candidate = RecognitionCandidate(
        expected_script=script,
        text=text,
        confidence=1.0,
        model_id="python-docx:native",
        provider_kind="docx_native",
    )
    winner, accepted, reasons = choose_best_recognition([candidate])
    # Native Word Unicode is not OCR. Unambiguous source-script spans therefore
    # do not inherit OCR-only confidence, quality, or length rejection rules.
    native_stats = script_statistics(winner.text)
    source_count = int(native_stats[script])
    if (
        source_count >= 1
        and winner.detected_script == script
        and winner.script_purity >= 0.90
    ):
        accepted = True
        reasons = []
    return LineResult(
        line=line_number,
        bbox=(0, line_number * 20, 600, line_number * 20 + 18),
        crop_file="",
        detection_confidence=1.0,
        detector="python-docx_native",
        script=winner.detected_script,
        text=winner.text,
        script_purity=winner.script_purity,
        confidence=winner.confidence,
        text_quality=winner.text_quality,
        recognition_model=winner.model_id,
        provider_kind=winner.provider_kind,
        accepted=accepted,
        review_required=not accepted,
        review_reasons=reasons,
        candidates=[_candidate_for_audit(enrich_candidate(candidate))],
    )


def _preserved_native_line_result(
    *,
    line_number: int,
    reference: str,
    text: str,
) -> LineResult:
    stats = script_statistics(text)
    has_letters = bool(int(stats["letters"]))
    return LineResult(
        line=line_number,
        bbox=(0, line_number * 20, 600, line_number * 20 + 18),
        crop_file="",
        detection_confidence=1.0,
        detector="python-docx_native",
        script="unknown",
        text=text,
        script_purity=float(stats["latin_ratio"]),
        confidence=1.0,
        text_quality=1.0,
        recognition_model="python-docx:native",
        provider_kind="docx_native",
        accepted=True,
        review_required=False,
        review_reasons=[],
        english=text,
        translation_status=(
            "preserved_english" if has_letters else "preserved_structured"
        ),
        candidates=[],
    )


def _paragraph_text_runs(paragraph: Any) -> list[Any]:
    """Return mutable text runs, including hyperlink-owned runs when available."""

    content_iterator = getattr(paragraph, "iter_inner_content", None)
    if not callable(content_iterator):
        return [run for run in paragraph.runs if run.text]
    runs: list[Any] = []
    for item in content_iterator():
        nested_runs = getattr(item, "runs", None)
        if nested_runs is not None:
            runs.extend(run for run in nested_runs if run.text)
        elif getattr(item, "text", ""):
            runs.append(item)
    return runs


def _replace_paragraph_text(paragraph: Any, replacement: str) -> bool:
    """Replace text across existing runs without removing drawings or styles."""

    runs = _paragraph_text_runs(paragraph)
    if not runs:
        if paragraph.text:
            return False
        paragraph.add_run(replacement)
        return True
    mutable_text = "".join(str(run.text or "") for run in runs)
    if normalize_text(mutable_text) != normalize_text(paragraph.text):
        # Text owned by an unsupported field/shape must not be duplicated or
        # destroyed. Preserve this paragraph and surface it for review.
        return False
    lengths = [max(1, len(str(run.text or ""))) for run in runs]
    total = sum(lengths)
    positions: list[int] = []
    cumulative = 0
    for length in lengths[:-1]:
        cumulative += length
        positions.append(round(len(replacement) * cumulative / total))
    values: list[str] = []
    start = 0
    for position in [*positions, len(replacement)]:
        values.append(replacement[start:position])
        start = position
    for run, value in zip(runs, values, strict=True):
        run.text = value
    return True


def _iter_image_blobs(document: Any) -> Iterator[tuple[str, bytes]]:
    seen: set[str] = set()
    for part in document.part.package.parts:
        content_type = str(getattr(part, "content_type", ""))
        if not content_type.startswith("image/"):
            continue
        blob = bytes(getattr(part, "blob", b""))
        digest = hashlib.sha256(blob).hexdigest()
        if not blob or digest in seen:
            continue
        seen.add(digest)
        yield digest[:12], blob


def _enhance_scanned_word_image(image: Image.Image) -> Image.Image:
    """Apply a fast, handwriting-safe grayscale enhancement for Tesseract."""

    source = ImageOps.exif_transpose(image).convert("RGB")
    rgb = np.asarray(source)
    if max(rgb.shape[:2]) > 2400:
        scale = 2400.0 / max(rgb.shape[:2])
        rgb = cv2.resize(rgb, None, fx=scale, fy=scale, interpolation=cv2.INTER_AREA)
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
    background = cv2.GaussianBlur(gray, (0, 0), sigmaX=24, sigmaY=24)
    normalized = cv2.divide(gray, np.maximum(background, 1), scale=235)
    clahe = cv2.createCLAHE(clipLimit=1.6, tileGridSize=(8, 8)).apply(normalized)
    # Small scans benefit from one conservative upscale; large scans avoid it.
    if max(clahe.shape) < 1400:
        clahe = cv2.resize(clahe, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
    return Image.fromarray(clahe).convert("RGB")


def _meaningful_latin_or_structured(text: str, confidence: float | None) -> bool:
    if confidence is None or confidence < 0.45:
        return False
    stripped = normalize_text(text)
    letters = sum(character.isalpha() for character in stripped)
    latin = int(script_statistics(stripped)["latin"])
    if letters >= 2 and latin / max(1, letters) >= 0.80:
        return True
    return bool(
        re.fullmatch(
            r"[\s\d.,:/()\-+%₹$#A-Za-z]{3,}",
            stripped,
        )
        and any(character.isdigit() for character in stripped)
    )


def _ocr_embedded_images(
    document: Any,
    output_directory: Path,
    progress_callback: ProgressCallback | None,
) -> list[LineResult]:
    """OCR each embedded image once with installed Punjabi/Hindi/English packs."""

    blobs = list(_iter_image_blobs(document))
    if not blobs:
        return []
    pytesseract, tessdata_option = _prepare_tesseract(("pan", "hin", "eng"))
    config = " ".join(
        item
        for item in (
            tessdata_option,
            "--oem 1 --psm 3 -c preserve_interword_spaces=1",
        )
        if item
    )
    crop_directory = output_directory / "embedded_line_crops"
    crop_directory.mkdir(parents=True, exist_ok=True)
    results: list[LineResult] = []
    for image_index, (digest, blob) in enumerate(blobs, start=1):
        _report_progress(
            progress_callback,
            f"Fast local OCR for embedded Word image {image_index}/{len(blobs)}",
        )
        try:
            with Image.open(BytesIO(blob)) as source:
                width, height = source.size
                if width < 240 or height < 120 or width * height < 60_000:
                    # Logos, icons, and signature thumbnails are preserved in
                    # the Word file and are not treated as page transcription.
                    continue
                enhanced = _enhance_scanned_word_image(source)
        except (UnidentifiedImageError, OSError):
            continue
        data = pytesseract.image_to_data(
            enhanced,
            lang="pan+hin+eng",
            config=config,
            output_type=pytesseract.Output.DICT,
        )
        groups: dict[tuple[int, int, int], list[int]] = {}
        count = len(data.get("text", []))
        for index in range(count):
            key = (
                int(data.get("block_num", [0] * count)[index] or 0),
                int(data.get("par_num", [0] * count)[index] or 0),
                int(data.get("line_num", [0] * count)[index] or 0),
            )
            groups.setdefault(key, []).append(index)
        for indices in groups.values():
            words: list[str] = []
            confidences: list[float] = []
            boxes: list[tuple[int, int, int, int]] = []
            for index in indices:
                word = normalize_text(str(data.get("text", [""] * count)[index]))
                try:
                    confidence = float(data.get("conf", [-1] * count)[index])
                except (TypeError, ValueError):
                    confidence = -1.0
                if not word:
                    continue
                words.append(word)
                if confidence >= 0:
                    confidences.append(confidence / 100.0)
                left = int(data.get("left", [0] * count)[index] or 0)
                top = int(data.get("top", [0] * count)[index] or 0)
                width = int(data.get("width", [0] * count)[index] or 0)
                height = int(data.get("height", [0] * count)[index] or 0)
                boxes.append((left, top, left + width, top + height))
            text = " ".join(words)
            if not text or not boxes:
                continue
            confidence_value = (
                sum(confidences) / len(confidences) if confidences else None
            )
            x1 = min(box[0] for box in boxes)
            y1 = min(box[1] for box in boxes)
            x2 = max(box[2] for box in boxes)
            y2 = max(box[3] for box in boxes)
            script = detect_script(text)
            line_number = len(results) + 1
            crop_name = f"embedded_line_crops/{digest}_{line_number:03d}.jpg"
            enhanced.crop((x1, y1, x2, y2)).save(
                output_directory / crop_name,
                quality=90,
            )
            if script in {"gurmukhi", "devanagari"}:
                candidate = RecognitionCandidate(
                    expected_script=script,
                    text=text,
                    confidence=confidence_value,
                    model_id="tesseract:pan+hin+eng",
                    provider_kind="tesseract_docx_scan",
                )
                winner, accepted, reasons = choose_best_recognition([candidate])
                result_script = winner.detected_script
                audit = [_candidate_for_audit(enrich_candidate(candidate))]
                text_quality = winner.text_quality
                purity = winner.script_purity
            else:
                accepted = _meaningful_latin_or_structured(text, confidence_value)
                reasons = [] if accepted else ["OCR output lacks reliable Indic or English text"]
                result_script = "unknown"
                text_quality = 1.0 if accepted else 0.0
                purity = float(script_statistics(text)["latin_ratio"])
                audit = []
            result = LineResult(
                line=line_number,
                bbox=(x1, y1, x2, y2),
                crop_file=crop_name,
                detection_confidence=confidence_value,
                detector="tesseract_docx_embedded_image",
                script=result_script,
                text=text,
                script_purity=purity,
                confidence=confidence_value,
                text_quality=text_quality,
                recognition_model="tesseract:pan+hin+eng",
                provider_kind="tesseract_docx_scan",
                accepted=accepted,
                review_required=not accepted,
                review_reasons=reasons,
                candidates=audit,
            )
            if accepted and script == "unknown":
                result.english = text
                result.translation_status = "preserved_english_or_structured"
            results.append(result)
    return results


def _refine_embedded_handwriting(
    lines: Sequence[LineResult],
    output_directory: Path,
    models: ModelRegistry,
    progress_callback: ProgressCallback | None,
) -> None:
    """Run neural HTR only for unresolved embedded-image scan lines."""

    targets = [line for line in lines if line.review_required and line.crop_file]
    if not targets:
        return
    crops: list[Image.Image] = []
    kept: list[LineResult] = []
    for line in targets:
        crop_path = output_directory / line.crop_file
        try:
            with Image.open(crop_path) as image:
                crops.append(image.convert("RGB"))
                kept.append(line)
        except (UnidentifiedImageError, OSError):
            continue
    if not crops:
        return
    _report_progress(
        progress_callback,
        f"Handwriting refinement for {len(crops)} unresolved Word scan lines",
    )
    gurmukhi = recognize_gurmukhi(crops, models, progress_callback, fast=False)
    unresolved: list[int] = []
    for index, candidate in enumerate(gurmukhi):
        winner, accepted, reasons = choose_best_recognition([candidate])
        line = kept[index]
        line.script = winner.detected_script
        line.text = winner.text
        line.script_purity = winner.script_purity
        line.confidence = winner.confidence
        line.text_quality = winner.text_quality
        line.recognition_model = winner.model_id
        line.provider_kind = winner.provider_kind
        line.accepted = accepted
        line.review_required = not accepted
        line.review_reasons = reasons
        line.candidates = [_candidate_for_audit(winner)]
        if not accepted:
            unresolved.append(index)
    if unresolved:
        hindi = recognize_hindi(
            [crops[index] for index in unresolved],
            models,
            progress_callback,
        )
        for index, candidate in zip(unresolved, hindi, strict=True):
            line = kept[index]
            winner, accepted, reasons = choose_best_recognition(
                [gurmukhi[index], candidate]
            )
            line.script = winner.detected_script
            line.text = winner.text
            line.script_purity = winner.script_purity
            line.confidence = winner.confidence
            line.text_quality = winner.text_quality
            line.recognition_model = winner.model_id
            line.provider_kind = winner.provider_kind
            line.accepted = accepted
            line.review_required = not accepted
            line.review_reasons = reasons
            line.candidates = [
                _candidate_for_audit(item) for item in (gurmukhi[index], candidate)
            ]
    models.release_recognizers()


def _render_native_unit(unit: _NativeUnit) -> str:
    values: list[str] = []
    for part in unit.parts:
        if part.line_result is None:
            values.append(part.text)
        elif (
            part.line_result.translation_status == "translated"
            and part.line_result.english
        ):
            values.append(part.line_result.english)
        else:
            # A failed/rejected translation is a no-op; source text remains.
            values.append(part.text)
    return "".join(values)


def _append_scanned_translation(document: Any, lines: Sequence[LineResult]) -> None:
    """Add selectable English after an image-only Word source without deleting it."""

    from docx.oxml.ns import qn
    from docx.shared import Pt

    if not lines:
        return
    document.add_page_break()
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run("English translation")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(15)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        if line.translation_status == "translated" and line.english:
            value = line.english
        elif line.translation_status == "preserved_english_or_structured":
            value = line.text
        else:
            value = f"[Line {line.line} requires review]"
        run = paragraph.add_run(value)
        run.font.name = "Aptos"
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")


def process_docx(
    input_path: Path,
    output_directory: Path,
    progress_callback: ProgressCallback | None = None,
    processing_mode: DocxProcessingMode = "fast_document",
) -> dict[str, Any]:
    """Transcribe and translate Word content through the fastest safe route."""

    if processing_mode not in {"fast_document", "line_accurate"}:
        raise ValueError(f"Unsupported DOCX processing mode: {processing_mode}")
    if input_path.suffix.casefold() != ".docx":
        raise ValueError("process_docx accepts only .docx input")
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency preflight handles this
        raise RuntimeError("python-docx is required for Word processing") from exc

    total_started = time.perf_counter()
    output_directory.mkdir(parents=True, exist_ok=True)
    _report_progress(progress_callback, "Reading native Word paragraphs and tables")
    try:
        document = Document(str(input_path))
    except Exception as exc:
        raise ValueError("The uploaded Word document is corrupted or unsupported") from exc

    extraction_started = time.perf_counter()
    native_units = _extract_native_units(document)
    native_indic = any(
        part.script is not None for unit in native_units for part in unit.parts
    )
    image_count = sum(1 for _item in _iter_image_blobs(document))
    lines: list[LineResult] = []
    for unit in native_units:
        translatable_parts = [
            part for part in unit.parts if part.script is not None and part.text.strip()
        ]
        if not translatable_parts:
            lines.append(
                _preserved_native_line_result(
                    line_number=len(lines) + 1,
                    reference=unit.reference,
                    text=unit.text,
                )
            )
            continue
        for part in unit.parts:
            if part.script is None or not part.text.strip():
                continue
            result = _native_line_result(
                line_number=len(lines) + 1,
                reference=unit.reference,
                text=part.text,
                script=part.script,
            )
            part.line_result = result
            lines.append(result)
    extraction_time = time.perf_counter() - extraction_started

    ocr_started = time.perf_counter()
    scan_lines: list[LineResult] = []
    if image_count:
        _report_progress(
            progress_callback,
            "Using fast local OCR on embedded Word images",
        )
        scan_lines = _ocr_embedded_images(
            document,
            output_directory,
            progress_callback,
        )
        line_offset = len(lines)
        for index, line in enumerate(scan_lines, start=1):
            line.line = line_offset + index
        lines.extend(scan_lines)
    ocr_time = time.perf_counter() - ocr_started

    models = ModelRegistry()
    if scan_lines and processing_mode == "line_accurate":
        _refine_embedded_handwriting(
            scan_lines,
            output_directory,
            models,
            progress_callback,
        )
    _report_progress(progress_callback, "Translating validated Word text to English")
    translation_time = translate_text(
        lines,
        models,
        progress_callback,
        num_beams=4 if processing_mode == "fast_document" else 6,
    )

    replacement_count = 0
    preserved_native_count = 0
    for unit in native_units:
        replacement = _render_native_unit(unit)
        if replacement == unit.text:
            continue
        if _replace_paragraph_text(unit.paragraph, replacement):
            replacement_count += 1
        else:
            preserved_native_count += 1
            for part in unit.parts:
                if part.line_result and part.line_result.translation_status == "translated":
                    part.line_result.translation_status = "preserved_unsafe_structure"
                    part.line_result.translation_error = (
                        "Word paragraph contains text not owned by safely mutable runs"
                    )
                    part.line_result.review_required = True
    if scan_lines:
        _append_scanned_translation(document, scan_lines)

    _report_progress(progress_callback, "Saving and validating translated Word document")
    translated_path = output_directory / "translated_en.docx"
    document.save(translated_path)
    reopened = Document(str(translated_path))
    if not reopened.sections:
        raise OSError("Translated Word output failed validation")

    transcription = "\n".join(
        line.text if line.accepted else f"[REVIEW REQUIRED - line {line.line}] {line.text}".rstrip()
        for line in lines
    )
    english_translation = "\n".join(
        line.english
        if line.translation_status
        in {
            "translated",
            "preserved_english",
            "preserved_structured",
            "preserved_english_or_structured",
        }
        else f"[REVIEW REQUIRED - line {line.line}]"
        for line in lines
    )
    (output_directory / "transcription.txt").write_text(
        transcription + ("\n" if transcription else ""), encoding="utf-8"
    )
    (output_directory / "translation_en.txt").write_text(
        english_translation + ("\n" if english_translation else ""), encoding="utf-8"
    )
    total_time = time.perf_counter() - total_started
    models.status["processing"] = {
        "mode": processing_mode,
        "native_text_bypassed_ocr": bool(native_units),
        "embedded_handwriting_refined": bool(
            scan_lines and processing_mode == "line_accurate"
        ),
    }
    timings = {
        "native_extraction": extraction_time,
        "embedded_image_ocr": ocr_time,
        "translation": translation_time,
        "total": total_time,
    }
    payload: dict[str, Any] = {
        "schema_version": "1.0",
        "input_document": str(input_path),
        "output_directory": str(output_directory),
        "document_type": "docx",
        "processing_route": (
            "native_docx_text" if native_indic else "embedded_image_tesseract" if scan_lines else "native_english_passthrough"
        ),
        "inference_only": True,
        "models": models.status,
        "native_paragraph_count": len(native_units),
        "embedded_image_count": image_count,
        "native_translation_replacement_count": replacement_count,
        "native_paragraphs_preserved_for_structure": preserved_native_count,
        "lines": [asdict(line) for line in lines],
        "transcription": transcription,
        "english_translation": english_translation,
        "primary_output": translated_path.name,
        "timings_seconds": timings,
        "accuracy_warning": (
            "Native Word text is exact Unicode extraction. Embedded-image OCR is "
            "probabilistic and uncertain lines require manual review."
        ),
    }
    (output_directory / "result.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return payload


__all__ = ["DocxProcessingMode", "process_docx", "split_translatable_parts"]
