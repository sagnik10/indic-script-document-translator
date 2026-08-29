"""Streamlit upload interface for the pretrained handwritten-page pipeline."""

from __future__ import annotations

import hashlib
import re
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from importlib.metadata import PackageNotFoundError, version
from io import BytesIO
from pathlib import Path
from typing import Any, Callable

import streamlit as st
from packaging.version import Version
from PIL import Image, UnidentifiedImageError


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from pretrained_page_ocr.kaggle_app import process_page  # noqa: E402

try:  # The document router is kept separate from the Streamlit view.
    from pretrained_page_ocr.document_processor import process_document  # noqa: E402
except ImportError:  # Allows image-only installations to keep starting cleanly.
    process_document = None


MAX_UPLOAD_BYTES = 100 * 1024 * 1024
MAX_IMAGE_PIXELS = 50_000_000
ALLOWED_FORMATS = {"JPEG": ".jpg", "PNG": ".png"}
ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".pdf", ".docx"}
DOWNLOAD_MIME_TYPES = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
    ".jpg": "image/jpeg",
    ".txt": "text/plain; charset=utf-8",
    ".json": "application/json",
    ".csv": "text/csv; charset=utf-8",
}


@dataclass(frozen=True, slots=True)
class UploadInfo:
    """Validated, display-safe metadata for one uploaded document."""

    kind: str
    description: str
    preview: bytes | None = None


def _safe_input_name(original_name: str, extension: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(original_name).stem).strip("._")
    return f"{stem or 'uploaded_document'}{extension}"


def _validate_image_upload(name: str, content: bytes) -> tuple[str, UploadInfo]:
    """Validate an image by decoded content rather than its extension alone."""

    try:
        with Image.open(BytesIO(content)) as image:
            image.verify()
        with Image.open(BytesIO(content)) as image:
            image_format = str(image.format or "").upper()
            dimensions = image.size
    except (UnidentifiedImageError, OSError) as exc:
        raise ValueError("The upload is not a valid JPG, JPEG, or PNG image.") from exc
    if image_format not in ALLOWED_FORMATS:
        raise ValueError(f"Unsupported image format: {image_format or 'unknown'}.")
    if dimensions[0] <= 0 or dimensions[1] <= 0:
        raise ValueError("The uploaded image has invalid dimensions.")
    if dimensions[0] * dimensions[1] > MAX_IMAGE_PIXELS:
        raise ValueError("The image is too large to process safely (maximum 50 megapixels).")
    return (
        _safe_input_name(name, ALLOWED_FORMATS[image_format]),
        UploadInfo(
            kind="image",
            description=f"{dimensions[0]} × {dimensions[1]} pixels",
            preview=content,
        ),
    )


def _validate_pdf_upload(name: str, content: bytes) -> tuple[str, UploadInfo]:
    """Validate PDF structure and render a lightweight first-page preview."""

    if not content.startswith(b"%PDF-"):
        raise ValueError("The upload does not contain a valid PDF signature.")
    try:
        import pymupdf

        with pymupdf.open(stream=content, filetype="pdf") as document:
            if document.needs_pass:
                raise ValueError("Password-protected PDFs are not supported.")
            if document.page_count < 1:
                raise ValueError("The PDF has no pages.")
            if document.page_count > 100:
                raise ValueError("The fast document interface supports up to 100 PDF pages per run.")
            page_count = document.page_count
            page = document.load_page(0)
            scale = min(
                1.25,
                1100.0 / max(1.0, page.rect.width),
                1500.0 / max(1.0, page.rect.height),
            )
            pixmap = page.get_pixmap(matrix=pymupdf.Matrix(scale, scale), alpha=False)
            preview = pixmap.tobytes("png")
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("The PDF is corrupted or cannot be opened.") from exc
    noun = "page" if page_count == 1 else "pages"
    return (
        _safe_input_name(name, ".pdf"),
        UploadInfo(kind="pdf", description=f"{page_count} {noun}", preview=preview),
    )


def _validate_docx_upload(name: str, content: bytes) -> tuple[str, UploadInfo]:
    """Validate that the ZIP payload is a genuine, readable Word document."""

    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            names = set(archive.namelist())
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                raise ValueError("The upload is not a valid DOCX document.")
            media_count = sum(name.startswith("word/media/") for name in names)

        from docx import Document

        document = Document(BytesIO(content))
        paragraph_count = sum(bool(paragraph.text.strip()) for paragraph in document.paragraphs)
        table_count = len(document.tables)
        if not paragraph_count and not table_count and not media_count:
            raise ValueError("The Word document is empty.")
    except ValueError:
        raise
    except (zipfile.BadZipFile, OSError, KeyError) as exc:
        raise ValueError("The DOCX file is corrupted or cannot be opened.") from exc
    description = (
        f"{paragraph_count} text paragraphs · {table_count} tables · "
        f"{media_count} embedded images"
    )
    return (
        _safe_input_name(name, ".docx"),
        UploadInfo(kind="docx", description=description),
    )


def _validate_upload(name: str, content: bytes) -> tuple[str, UploadInfo]:
    if not content:
        raise ValueError("The uploaded file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError("The document exceeds the 100 MB upload limit.")
    extension = Path(name).suffix.casefold()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Supported formats are PDF, DOCX, JPG, JPEG, and PNG.")
    if extension == ".pdf":
        return _validate_pdf_upload(name, content)
    if extension == ".docx":
        return _validate_docx_upload(name, content)
    return _validate_image_upload(name, content)


def _process_upload(
    content: bytes,
    safe_name: str,
    content_digest: str,
    processing_mode: str,
    progress_callback: Callable[[str], None] | None = None,
) -> tuple[dict[str, Any], dict[str, bytes]]:
    """Process one immutable upload and return serializable output artifacts."""

    del content_digest  # Included in the cache key for explicit provenance.
    with tempfile.TemporaryDirectory(prefix="indic_htr_ui_") as temp_name:
        temp_directory = Path(temp_name)
        input_path = temp_directory / safe_name
        output_directory = temp_directory / "page_ocr_result"
        input_path.write_bytes(content)
        if process_document is not None:
            payload = process_document(
                input_path,
                output_directory,
                progress_callback=progress_callback,
                processing_mode=processing_mode,
            )
        elif input_path.suffix.casefold() in {".jpg", ".jpeg", ".png"}:
            payload = process_page(
                input_path,
                output_directory,
                progress_callback=progress_callback,
                processing_mode=(
                    "line_accurate" if processing_mode == "line_accurate" else "fast_cpu"
                ),
            )
        else:
            raise RuntimeError(
                "The fast PDF/DOCX document processor is not available in this installation."
            )
        artifacts = {
            path.relative_to(output_directory).as_posix(): path.read_bytes()
            for path in output_directory.rglob("*")
            if path.is_file()
        }
    return payload, artifacts


def _download_name(upload_name: str, artifact_name: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(upload_name).stem).strip("._")
    artifact = Path(artifact_name)
    return f"{stem or 'document'}_{artifact.stem}{artifact.suffix}"


def _result_rows(payload: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for line in payload.get("lines", []):
        rows.append(
            {
                "Page": line.get("page_number", 1),
                "Line": line.get("line"),
                "Script": line.get("script"),
                "Recognized text": line.get("text"),
                "OCR confidence": line.get("confidence"),
                "Script purity": line.get("script_purity"),
                "Accepted": line.get("accepted"),
                "English": line.get("english"),
                "Translation confidence": line.get("translation_confidence"),
                "Translation quality": line.get("translation_quality"),
                "Context group": line.get("translation_context_group"),
                "Translation status": line.get("translation_status"),
                "Review reason": "; ".join(
                    [
                        *line.get("review_reasons", []),
                        *line.get("translation_review_reasons", []),
                    ]
                ),
            }
        )
    return rows


def _runtime_dependency_error() -> str | None:
    try:
        version("PyMuPDF")
    except PackageNotFoundError:
        return "PyMuPDF is not installed; PDF documents cannot be opened."
    try:
        version("python-docx")
    except PackageNotFoundError:
        return "python-docx is not installed; the translated Word document cannot be created."
    try:
        Version(version("torchvision"))
    except (PackageNotFoundError, ValueError):
        return "Torchvision is not installed; the pretrained OCR vision models cannot load."
    try:
        installed = Version(version("transformers"))
    except (PackageNotFoundError, ValueError):
        return "Transformers is not installed."
    required = Version("5.12.1")
    if installed < required:
        return f"Transformers {installed} is installed, but this app requires {required} or newer."
    return None


st.set_page_config(
    page_title="Hindi and Gurmukhi document translator",
    page_icon="📝",
    layout="wide",
)

st.session_state.setdefault("pretrained_ocr_result", None)
st.session_state.setdefault("pretrained_ocr_digest", None)

st.title("Hindi and Gurmukhi document translator")
st.caption(
    "Upload a PDF, Word document, scan, or photograph. Native PDF/DOCX text is "
    "extracted directly; OCR runs only where a page is actually scanned."
)
st.info(
    "Fast document mode avoids handwriting models when native text is available. "
    "Uncertain text is preserved for review rather than guessed.",
    icon=":material/info:",
)

cuda_available = False
try:
    import torch

    cuda_available = torch.cuda.is_available()
    if not cuda_available:
        st.warning(
            "CUDA is not available. Native PDF/DOCX extraction is fast on CPU. "
            "Choose handwriting recognition only for genuinely handwritten sources.",
            icon=":material/speed:",
        )
except ImportError:
    pass

dependency_error = _runtime_dependency_error()
if dependency_error:
    st.error(dependency_error, icon=":material/package_2:")
    st.code(
        r".\.venv\Scripts\python.exe -m pip install -r "
        r".\pretrained_page_ocr\requirements-ui-windows.txt",
        language="powershell",
    )
    st.stop()

uploaded_file = st.file_uploader(
    "Upload a document",
    type=["pdf", "docx", "jpg", "jpeg", "png"],
    max_upload_size=100,
    help="Supported formats: PDF, DOCX, JPG, JPEG, and PNG. Maximum size: 100 MB.",
    key="pretrained_page_upload",
)

if uploaded_file is None:
    st.session_state.pretrained_ocr_result = None
    st.session_state.pretrained_ocr_digest = None
    st.stop()

upload_bytes = uploaded_file.getvalue()
try:
    safe_filename, upload_info = _validate_upload(uploaded_file.name, upload_bytes)
except ValueError as exc:
    st.error(str(exc), icon=":material/error:")
    st.stop()

digest = hashlib.sha256(upload_bytes).hexdigest()

with st.container(border=True):
    st.markdown(f"**{uploaded_file.name}**")
    st.caption(f"{upload_info.kind.upper()} · {upload_info.description}")
    if upload_info.preview is not None:
        st.image(upload_info.preview, caption="Source preview")

with st.form("recognize_page", border=False):
    processing_mode = st.segmented_control(
        "Processing mode",
        options=["fast_document", "line_accurate"],
        default="fast_document",
        format_func=lambda value: {
            "fast_document": "Fast document",
            "line_accurate": "Hybrid handwriting accuracy",
        }[value],
        help=(
            "Fast document extracts native PDF/DOCX text first and uses Tesseract on "
            "scanned pages. Hybrid mode keeps that fast route, then sends only unresolved "
            "handwritten lines to the slower neural recognizers. Both modes now use "
            "validated contextual translation and quality beam search; Hybrid uses the "
            "strongest decoding and is slower."
        ),
        width="stretch",
    )
    submitted = st.form_submit_button(
        "Transcribe and translate document",
        type="primary",
        icon=":material/translate:",
        width="stretch",
    )

request_identity = f"{digest}:{processing_mode}"
if st.session_state.pretrained_ocr_digest != request_identity:
    st.session_state.pretrained_ocr_result = None
    st.session_state.pretrained_ocr_digest = request_identity

if submitted:
    status = st.status("Preparing the page and loading pretrained models…", expanded=True)
    try:
        def report_progress(message: str) -> None:
            status.update(label=message, state="running", expanded=True)
            status.write(message)

        report_progress("Extracting document text and OCRing only the pages that need it")
        payload, artifacts = _process_upload(
            upload_bytes,
            safe_filename,
            digest,
            processing_mode or "fast_document",
            progress_callback=report_progress,
        )
        st.session_state.pretrained_ocr_result = {
            "payload": payload,
            "artifacts": artifacts,
            "upload_name": uploaded_file.name,
            "identity": request_identity,
        }
        status.update(label="Recognition and translation complete", state="complete", expanded=False)
    except Exception as exc:
        status.update(label="Processing failed", state="error", expanded=True)
        st.error(
            f"Processing could not complete: {type(exc).__name__}: {exc}",
            icon=":material/error:",
        )

result = st.session_state.pretrained_ocr_result
if result is None:
    st.stop()

payload = result["payload"]
artifacts = result["artifacts"]
lines = payload.get("lines", [])
accepted_count = sum(bool(line.get("accepted")) for line in lines)
translated_count = sum(line.get("translation_status") == "translated" for line in lines)
review_count = sum(bool(line.get("review_required")) for line in lines)
translated_quality = [
    float(line.get("translation_quality") or 0.0)
    for line in lines
    if line.get("translation_status") == "translated"
]

st.subheader("Results")
with st.container(horizontal=True):
    st.metric("Detected lines", len(lines))
    st.metric("Accepted lines", accepted_count)
    st.metric("Translated lines", translated_count)
    st.metric("Review required", review_count)
    st.metric(
        "Translation QA",
        f"{sum(translated_quality) / len(translated_quality):.0%}"
        if translated_quality
        else "—",
    )

st.caption(
    "Translation QA checks English output shape, repetition, and exact preservation of "
    "dates, identifiers, numbers, and existing English. It is not a certified accuracy "
    "percentage; OCR source errors still require review."
)

if "detected_lines.jpg" in artifacts:
    with st.container(border=True):
        st.image(artifacts["detected_lines.jpg"], caption="Detected lines: green is accepted; orange needs review")

st.text_area(
    "Punjabi/Hindi transcription",
    value=payload.get("transcription", ""),
    height=220,
    disabled=True,
    key=f"transcription_{result.get('identity', digest)}",
)
st.text_area(
    "English translation",
    value=payload.get("english_translation", ""),
    height=220,
    disabled=True,
    key=f"translation_{result.get('identity', digest)}",
)

rows = _result_rows(payload)
if rows:
    st.subheader("Line review")
    st.dataframe(rows, hide_index=True)

st.subheader("Translated Word document")
word_artifact = "translated_en.docx"
word_content = artifacts.get(word_artifact)
if word_content is not None:
    st.download_button(
        label="Download translated Word document",
        data=word_content,
        file_name=_download_name(result["upload_name"], word_artifact),
        mime=DOWNLOAD_MIME_TYPES[".docx"],
        icon=":material/download:",
        type="primary",
        width="stretch",
        on_click="ignore",
        key=f"download_{result.get('identity', digest)}_{word_artifact}",
    )
    st.caption(
        f"{translated_count} of {len(lines)} detected lines translated. "
        "Every other detected line is retained as an explicit review entry; no text is invented."
    )
else:
    st.error("The translated Word artifact was not generated.", icon=":material/error:")

with st.expander("Additional downloads"):
    preferred_order = [
        "translation_en.txt",
        "transcription.txt",
        "result.json",
        "result.csv",
        "detected_lines.jpg",
        "corrected_page.jpg",
    ]
    for artifact_name in preferred_order:
        content = artifacts.get(artifact_name)
        if content is None:
            continue
        suffix = Path(artifact_name).suffix.casefold()
        st.download_button(
            label=f"Download {artifact_name}",
            data=content,
            file_name=_download_name(result["upload_name"], artifact_name),
            mime=DOWNLOAD_MIME_TYPES.get(suffix, "application/octet-stream"),
            icon=":material/download:",
            on_click="ignore",
            key=f"download_{result.get('identity', digest)}_{artifact_name}",
        )

with st.expander("Model and timing details"):
    st.json(
        {
            "models": payload.get("models", {}),
            "timings_seconds": payload.get("timings_seconds", {}),
            "preprocessing_operations": payload.get("preprocessing_operations", []),
            "accuracy_warning": payload.get("accuracy_warning"),
        },
        expanded=False,
    )
