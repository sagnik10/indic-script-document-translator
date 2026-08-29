"""Reopen generated artifacts and validate basic structural integrity."""

from __future__ import annotations

import io

from ..exceptions import RenderingError
from ..schemas import DocumentModel, TranslationStatus


def validate_pdf_output(data: bytes, source: DocumentModel) -> list[str]:
    if not data.startswith(b"%PDF-"):
        raise RenderingError("Generated PDF signature is invalid")
    try:
        import pymupdf as fitz

        with fitz.open(stream=data, filetype="pdf") as pdf:
            if pdf.page_count != len(source.pages):
                raise RenderingError("Generated PDF page count differs from the source")
            for page_model in source.pages:
                page = pdf.load_page(page_model.page_number - 1)
                if (
                    abs(page.rect.width - page_model.width) > 0.05
                    or abs(page.rect.height - page_model.height) > 0.05
                    or int(page.rotation) != int(page_model.rotation)
                ):
                    raise RenderingError("Generated PDF page geometry differs from the source")
            rendered = [
                block
                for block in source.blocks
                if block.translation_status == TranslationStatus.TRANSLATED
                and block.metadata.get("replacement_applied") is True
            ]
            if rendered and not any(
                pdf.load_page(index).get_text().strip() for index in range(pdf.page_count)
            ):
                raise RenderingError("Generated PDF contains no searchable translated text")
            for block in rendered:
                page_text = " ".join(
                    pdf.load_page(block.page_number - 1).get_text().split()
                ).casefold()
                expected = " ".join((block.english_translation or "").split()).casefold()
                if expected and expected not in page_text:
                    raise RenderingError(
                        "A rendered translation is not searchable on its source page"
                    )
            if source.metadata.get("diagnostic_pages_appended", 0):
                raise RenderingError("Diagnostic pages were appended to the primary PDF")
            if source.metadata.get("primary_output_has_debug_overlays"):
                raise RenderingError("Debug overlays were included in the primary PDF")
    except ImportError:
        return ["PyMuPDF was unavailable for post-render PDF validation."]
    except RenderingError:
        raise
    except Exception as exc:
        raise RenderingError("Generated PDF could not be reopened") from exc
    return []


def validate_docx_output(data: bytes, source: DocumentModel) -> list[str]:
    if not data.startswith(b"PK"):
        raise RenderingError("Generated DOCX container signature is invalid")
    try:
        from docx import Document

        document = Document(io.BytesIO(data))
        has_content = any(paragraph.text.strip() for paragraph in document.paragraphs)
        has_content = has_content or bool(document.tables) or bool(document.inline_shapes)
        if not has_content:
            raise RenderingError("Generated DOCX is empty")
    except ImportError:
        return ["python-docx was unavailable for post-render DOCX validation."]
    except RenderingError:
        raise
    except Exception as exc:
        raise RenderingError("Generated DOCX could not be reopened") from exc
    return []


def validate_output(data: bytes, extension: str, source: DocumentModel) -> list[str]:
    if not data:
        raise RenderingError("Generated output is empty")
    if extension == "pdf":
        return validate_pdf_output(data, source)
    if extension == "docx":
        return validate_docx_output(data, source)
    raise RenderingError(f"Unsupported output extension: {extension}")
