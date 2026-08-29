"""Kaggle-ready, inference-only Hindi/Gurmukhi handwritten page OCR.

Change ``PAGE_IMAGE`` below, run this file on a Kaggle GPU, and collect the
artifacts written to ``OUTPUT_DIRECTORY``.  The program never trains or writes
model weights.  Hugging Face uses its normal read-through model cache.

Kaggle install cell (run once before this script)::

    !sudo apt-get update -qq && sudo apt-get install -y -qq tesseract-ocr tesseract-ocr-pan tesseract-ocr-hin
    %pip install -q "transformers>=5.12.1,<6" "surya-ocr==0.22.1" \
        "pillow>=10.2,<11" "opencv-python-headless==4.11.0.86" \
        "sentencepiece>=0.2,<1" "protobuf>=5,<7" "pytesseract>=0.3.13,<1" \
        "python-docx>=1.1,<2"

The primary Hindi recognizer is the requested Devanagari TrOCR checkpoint.
No credible public *line-level handwritten* Gurmukhi checkpoint was found.
The Gurmukhi path therefore uses Surya OCR 2 as a clearly reported
multilingual fallback.  Its public documentation includes handwritten-note
examples and Punjabi in its multilingual benchmark.  Official Tesseract
``pan`` is a last-resort fallback if Surya cannot load; it is not claimed to be
a handwriting-specialized model.
"""

from __future__ import annotations

import argparse
import contextlib
import csv
import gc
import html
import importlib
import json
import math
import os
import re
import shutil
import statistics
import sys
import time
import unicodedata
from dataclasses import asdict, dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Iterable, Iterator, Literal, Sequence

import cv2
import numpy as np
from PIL import Image, ImageOps


# ---------------------------------------------------------------------------
# User configuration: this is normally the only value that needs changing.
# ---------------------------------------------------------------------------
PAGE_IMAGE: str | None = "/kaggle/input/.../test.jpg"

OUTPUT_DIRECTORY = Path("/kaggle/working/page_ocr_result")
HINDI_MODEL_ID = "paudelanil/trocr-devanagari-2"
GURMUKHI_MODEL_ID = "datalab-to/surya-ocr-2"
TRANSLATION_MODEL_ID = "facebook/nllb-200-distilled-600M"

OCR_BATCH_SIZE = 8
VLM_BATCH_SIZE = 4
TRANSLATION_BATCH_SIZE = 8
TRANSLATION_CONTEXT_MAX_LINES = 3
TRANSLATION_CONTEXT_MAX_CHARACTERS = 420
LINE_PADDING_PIXELS = 12
MIN_LINE_WIDTH = 28
MIN_LINE_HEIGHT = 8
MIN_LINE_AREA = 260
MIN_SOURCE_LETTERS = 2
MIN_AUTOMATIC_LINE_LETTERS = 4
MIN_SCRIPT_PURITY = 0.62
MIN_OCR_CONFIDENCE = 0.28
MIN_TEXT_QUALITY = 0.52
MAX_IMAGE_SIDE = 2600

ScriptName = Literal["gurmukhi", "devanagari", "mixed", "unknown"]
BBox = tuple[int, int, int, int]
ProgressCallback = Callable[[str], None]
ProcessingMode = Literal["fast_cpu", "line_accurate"]

GURMUKHI_START, GURMUKHI_END = 0x0A00, 0x0A7F
DEVANAGARI_START, DEVANAGARI_END = 0x0900, 0x097F


def _from_pretrained_cached_first(factory: Any, model_id: str, **kwargs: Any) -> Any:
    """Use a complete local Hugging Face cache before making a Hub request."""

    try:
        return factory.from_pretrained(
            model_id,
            local_files_only=True,
            **kwargs,
        )
    except (OSError, RuntimeError):
        return factory.from_pretrained(model_id, **kwargs)


def _use_max_new_tokens_only(model: Any) -> None:
    """Prevent Transformers' max_length/max_new_tokens conflict warning."""

    generation_config = getattr(model, "generation_config", None)
    if generation_config is not None:
        generation_config.max_length = None


def _finish_pretrained_model_load(model: Any, device: Any) -> Any:
    """Resolve tied weights, reject meta tensors, and then move a real model.

    ``to_empty`` is intentionally never used here: it would materialize random
    parameters and make OCR appear operational with invalid weights.
    """

    tie_weights = getattr(model, "tie_weights", None)
    if callable(tie_weights):
        tie_weights()
    tensors: list[tuple[str, Any]] = []
    named_parameters = getattr(model, "named_parameters", None)
    if callable(named_parameters):
        tensors.extend((f"parameter:{name}", value) for name, value in named_parameters())
    named_buffers = getattr(model, "named_buffers", None)
    if callable(named_buffers):
        tensors.extend((f"buffer:{name}", value) for name, value in named_buffers())
    meta_names = [name for name, value in tensors if getattr(value, "is_meta", False)]
    if meta_names:
        preview = ", ".join(meta_names[:4])
        raise RuntimeError(
            "The pretrained checkpoint left unmaterialized meta tensors "
            f"({preview}). The model was rejected instead of initializing fake weights."
        )
    target = str(device)
    needs_move = not tensors or any(str(value.device) != target for _name, value in tensors)
    if needs_move:
        model = model.to(device)
    remaining_meta = [
        name
        for name, value in (
            list(model.named_parameters()) + list(model.named_buffers())
        )
        if getattr(value, "is_meta", False)
    ] if callable(getattr(model, "named_parameters", None)) and callable(
        getattr(model, "named_buffers", None)
    ) else []
    if remaining_meta:
        raise RuntimeError("The pretrained model still contains meta tensors after loading")
    return model


def _find_tesseract_executable() -> Path | None:
    """Find Tesseract from PATH, an explicit override, or standard Windows paths."""

    candidates: list[str | Path | None] = [
        os.environ.get("TESSERACT_CMD"),
        shutil.which("tesseract"),
        Path(os.environ.get("ProgramFiles", r"C:\Program Files"))
        / "Tesseract-OCR"
        / "tesseract.exe",
        Path(os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"))
        / "Tesseract-OCR"
        / "tesseract.exe",
    ]
    for candidate in candidates:
        if not candidate:
            continue
        path = Path(candidate).expanduser()
        if path.is_file():
            return path.resolve()
    return None


def _find_tessdata_directory(required_languages: Sequence[str]) -> Path | None:
    """Locate a tessdata directory containing every required language pack."""

    executable = _find_tesseract_executable()
    candidates: list[str | Path | None] = [
        os.environ.get("TESSDATA_PREFIX"),
        Path(__file__).resolve().parents[1] / ".runtime" / "tessdata",
        executable.parent / "tessdata" if executable else None,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        directory = Path(candidate).expanduser()
        if directory.name.casefold() != "tessdata" and (directory / "tessdata").is_dir():
            directory = directory / "tessdata"
        if directory.is_dir() and all(
            (directory / f"{language}.traineddata").is_file()
            for language in required_languages
        ):
            return directory.resolve()
    return None


def _prepare_tesseract(required_languages: Sequence[str]) -> tuple[Any, str]:
    """Configure pytesseract for local Windows installs and bundled language data."""

    import pytesseract

    executable = _find_tesseract_executable()
    if executable is not None:
        pytesseract.pytesseract.tesseract_cmd = str(executable)
    tessdata = _find_tessdata_directory(required_languages)
    if tessdata is not None:
        # Passing a quoted ``--tessdata-dir`` through pytesseract leaves literal
        # quote characters in the Windows path. TESSDATA_PREFIX avoids that
        # subprocess parsing problem and is also supported by native Tesseract.
        os.environ["TESSDATA_PREFIX"] = str(tessdata)
    return pytesseract, ""


@dataclass(slots=True)
class EnhancedPage:
    """Image variants kept separate so faint strokes are never discarded."""

    corrected_rgb: np.ndarray
    enhanced_rgb: np.ndarray
    enhanced_gray: np.ndarray
    threshold: np.ndarray
    operations: list[str] = field(default_factory=list)


@dataclass(slots=True)
class TextRegion:
    """A detected logical text line in corrected-page coordinates."""

    bbox: BBox
    detection_confidence: float | None
    detector: str


@dataclass(slots=True)
class RecognitionCandidate:
    """One model's source-language reading of a line crop."""

    expected_script: ScriptName
    text: str
    confidence: float | None
    model_id: str
    provider_kind: str
    error: str | None = None
    detected_script: ScriptName = "unknown"
    script_purity: float = 0.0
    text_quality: float = 0.0
    source_letter_count: int = 0
    score: float = 0.0


@dataclass(slots=True)
class LineResult:
    """Auditable recognition, routing, and translation result for one line."""

    line: int
    bbox: BBox
    crop_file: str
    detection_confidence: float | None
    detector: str
    script: ScriptName
    text: str
    script_purity: float
    confidence: float | None
    text_quality: float
    recognition_model: str
    provider_kind: str
    accepted: bool
    review_required: bool
    review_reasons: list[str]
    english: str = ""
    translation_status: str = "not_attempted"
    translation_error: str | None = None
    translation_confidence: float | None = None
    translation_quality: float = 0.0
    translation_review_reasons: list[str] = field(default_factory=list)
    translation_context_group: str | None = None
    candidates: list[dict[str, Any]] = field(default_factory=list)
    page_number: int = 1


@dataclass(frozen=True, slots=True)
class ProtectedTranslationSource:
    """A source string with exact values hidden from machine translation."""

    text: str
    replacements: dict[str, str]
    source_tokens: tuple[str, ...]


@dataclass(slots=True)
class TranslationUnit:
    """One contextual source unit mapped back to one or more physical lines."""

    lines: list[LineResult]
    script: ScriptName
    source: str
    protected: ProtectedTranslationSource
    group_id: str


TRANSLATION_PROTECTED_PATTERN = re.compile(
    r"(?:https?://[^\s]+|www\.[^\s]+|[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}|"
    r"(?:\+?\d[\d\s().-]{6,}\d)|(?:₹|Rs\.?|INR|\$|€|£)\s?\d[\d,]*(?:\.\d+)?|"
    r"\b\d{1,4}[-/.]\d{1,2}[-/.]\d{1,4}\b|"
    r"\b(?:Sec(?:tion)?\.?|FIR|Case|File|Ref(?:erence)?|Dispatch|MRD|OPD|IPD)\s*"
    r"(?:No\.?\s*)?[A-Za-z0-9_./()-]+|"
    r"\b\d+(?:[.,:/-]\d+)*%?\b|"
    # Existing English names, institutions, abbreviations, and medical terms
    # are pass-through content on an otherwise Indic source line.
    r"\b[A-Za-z][A-Za-z0-9&.'()/+-]*(?:\s+[A-Za-z][A-Za-z0-9&.'()/+-]*)*\b)",
    re.IGNORECASE,
)


def normalize_text(value: str) -> str:
    """Normalize Unicode and whitespace without guessing source content."""

    value = html.unescape(value or "")
    value = re.sub(r"<[^>]+>", " ", value)
    value = unicodedata.normalize("NFC", value)
    value = value.replace("\u200b", "").replace("\ufeff", "")
    return " ".join(value.strip().split())


def normalize_translation_source(value: str) -> str:
    """Normalize OCR spacing that is unambiguously invalid inside Indic glyphs.

    This only joins combining signs and virama sequences to their base glyph. It
    deliberately does not guess whether two ordinary source letters form one
    word, so it cannot silently reconstruct missing source content.
    """

    normalized = normalize_text(value)
    output: list[str] = []
    pending_space = False
    viramas = {"\u094d", "\u0a4d"}
    for character in normalized:
        if character.isspace():
            pending_space = bool(output)
            continue
        category = unicodedata.category(character)
        attaches_to_previous = category.startswith("M")
        follows_virama = bool(output and output[-1] in viramas)
        if pending_space and not attaches_to_previous and not follows_virama:
            output.append(" ")
        output.append(character)
        pending_space = False
    result = "".join(output)
    result = re.sub(r"\s+([।॥,.;:!?])", r"\1", result)
    return normalize_text(result)


def _translation_placeholder(index: int) -> str:
    value = index
    letters = ""
    while True:
        letters = chr(ord("A") + value % 26) + letters
        value = value // 26 - 1
        if value < 0:
            break
    return f"ZXQPROTECTED{letters}QXZ"


def protect_translation_source(value: str) -> ProtectedTranslationSource:
    """Protect structured values and existing Latin text before NLLB."""

    replacements: dict[str, str] = {}
    tokens: list[str] = []

    def replace(match: re.Match[str]) -> str:
        placeholder = _translation_placeholder(len(replacements))
        replacements[placeholder] = match.group(0)
        tokens.append(match.group(0))
        return placeholder

    protected = TRANSLATION_PROTECTED_PATTERN.sub(replace, value)
    return ProtectedTranslationSource(
        text=protected,
        replacements=replacements,
        source_tokens=tuple(tokens),
    )


def restore_translation_source(
    value: str,
    protected: ProtectedTranslationSource,
) -> tuple[str, list[str]]:
    """Restore protected values and report any placeholder the model lost."""

    restored = value
    missing: list[str] = []
    for placeholder, source_token in protected.replacements.items():
        # Generative tokenizers sometimes add spaces or punctuation around an
        # otherwise unchanged Latin sentinel. Match those harmless variations.
        characters = [re.escape(character) for character in placeholder]
        flexible = re.compile(r"[\s_\-]*".join(characters), re.IGNORECASE)
        restored, count = flexible.subn(lambda _match, token=source_token: token, restored)
        if count == 0:
            missing.append(source_token)
    return normalize_text(restored), missing


def translation_quality(
    source: str,
    translated: str,
    missing_protected_tokens: Sequence[str] = (),
) -> tuple[float, list[str]]:
    """Validate that NLLB returned bounded English rather than corrupt output."""

    output = normalize_text(translated)
    reasons: list[str] = []
    if not output:
        return 0.0, ["translation model returned empty text"]
    stats = script_statistics(output)
    letters = int(stats["letters"])
    latin = int(stats["latin"])
    latin_ratio = latin / max(1, letters)
    meaningful_words = re.findall(r"[A-Za-z][A-Za-z'-]{1,}", output)
    if not meaningful_words:
        reasons.append("translation has no meaningful English words")
    if letters and latin_ratio < 0.65:
        reasons.append("translation output is not predominantly English")
    source_letters = max(1, int(script_statistics(source)["letters"]))
    length_ratio = len(output) / source_letters
    if length_ratio < 0.18 or length_ratio > 8.0:
        reasons.append("translation length is implausible for the validated source")
    words = re.findall(r"\w+", output.casefold())
    repeated_ratio = 0.0
    if words:
        repeated_ratio = 1.0 - len(set(words)) / len(words)
    if len(words) >= 8 and repeated_ratio > 0.62:
        reasons.append("translation contains excessive repetition")
    if missing_protected_tokens:
        reasons.append("translation did not preserve every protected value")

    length_score = max(0.0, 1.0 - abs(min(3.5, length_ratio) - 1.25) / 3.5)
    score = (
        0.45 * min(1.0, latin_ratio)
        + 0.25 * length_score
        + 0.20 * (1.0 - min(1.0, repeated_ratio))
        + 0.10 * (0.0 if missing_protected_tokens else 1.0)
    )
    if reasons:
        score *= 0.55
    return float(max(0.0, min(1.0, score))), reasons


def script_statistics(text: str) -> dict[str, float | int]:
    """Return Unicode-script counts and ratios over linguistic characters."""

    normalized = normalize_text(text)
    letters = [character for character in normalized if character.isalpha()]
    gurmukhi = sum(
        GURMUKHI_START <= ord(character) <= GURMUKHI_END for character in letters
    )
    devanagari = sum(
        DEVANAGARI_START <= ord(character) <= DEVANAGARI_END for character in letters
    )
    latin = sum("LATIN" in unicodedata.name(character, "") for character in letters)
    total = len(letters)
    return {
        "letters": total,
        "gurmukhi": gurmukhi,
        "devanagari": devanagari,
        "latin": latin,
        "gurmukhi_ratio": gurmukhi / max(1, total),
        "devanagari_ratio": devanagari / max(1, total),
        "latin_ratio": latin / max(1, total),
    }


def detect_script(text: str) -> ScriptName:
    """Classify recognized Unicode independently from the selected model."""

    stats = script_statistics(text)
    gurmukhi = int(stats["gurmukhi"])
    devanagari = int(stats["devanagari"])
    if gurmukhi and devanagari:
        return "mixed"
    if gurmukhi:
        return "gurmukhi"
    if devanagari:
        return "devanagari"
    return "unknown"


def script_purity(text: str, script: ScriptName) -> float:
    stats = script_statistics(text)
    if script == "gurmukhi":
        return float(stats["gurmukhi_ratio"])
    if script == "devanagari":
        return float(stats["devanagari_ratio"])
    if script == "mixed":
        return float(stats["gurmukhi_ratio"]) + float(stats["devanagari_ratio"])
    return 0.0


def calculate_text_quality(text: str, expected_script: ScriptName) -> float:
    """Conservatively score whether OCR output resembles source-language text."""

    normalized = normalize_text(text)
    if not normalized:
        return 0.0
    stats = script_statistics(normalized)
    letters = int(stats["letters"])
    purity = script_purity(normalized, expected_script)
    visible = [character for character in normalized if not character.isspace()]
    symbol_ratio = sum(
        not character.isalnum() and not character.isspace() for character in visible
    ) / max(1, len(visible))
    replacement_markers = {"?", "\ufffd", "□", "■", "|"}
    replacement_ratio = sum(character in replacement_markers for character in visible) / max(
        1, len(visible)
    )
    repeated = bool(re.search(r"(.)\1{3,}", normalized))
    length_score = min(1.0, letters / 8.0)
    quality = (
        0.55 * purity
        + 0.25 * length_score
        + 0.20 * max(0.0, 1.0 - symbol_ratio - replacement_ratio)
    )
    if repeated:
        quality *= 0.55
    if letters < MIN_SOURCE_LETTERS:
        quality *= 0.45
    return float(max(0.0, min(1.0, quality)))


def _candidate_score(candidate: RecognitionCandidate) -> float:
    confidence = candidate.confidence if candidate.confidence is not None else 0.35
    fallback_penalty = 0.04 if candidate.provider_kind == "tesseract_fallback" else 0.0
    return max(
        0.0,
        0.48 * candidate.script_purity
        + 0.30 * candidate.text_quality
        + 0.22 * confidence
        - fallback_penalty,
    )


def enrich_candidate(candidate: RecognitionCandidate) -> RecognitionCandidate:
    candidate.text = normalize_text(candidate.text)
    candidate.detected_script = detect_script(candidate.text)
    candidate.script_purity = script_purity(candidate.text, candidate.expected_script)
    candidate.text_quality = calculate_text_quality(
        candidate.text, candidate.expected_script
    )
    stats = script_statistics(candidate.text)
    key = "gurmukhi" if candidate.expected_script == "gurmukhi" else "devanagari"
    candidate.source_letter_count = int(stats[key])
    candidate.score = _candidate_score(candidate)
    return candidate


def choose_best_recognition(
    candidates: Sequence[RecognitionCandidate],
) -> tuple[RecognitionCandidate, bool, list[str]]:
    """Choose script-consistent OCR and surface uncertainty rather than hiding it."""

    if not candidates:
        empty = RecognitionCandidate(
            expected_script="unknown",
            text="",
            confidence=None,
            model_id="none",
            provider_kind="unavailable",
            error="no recognition candidate was produced",
        )
        return empty, False, ["no recognition candidate"]

    enriched = [enrich_candidate(candidate) for candidate in candidates]
    ordered = sorted(enriched, key=lambda candidate: candidate.score, reverse=True)
    winner = ordered[0]
    reasons: list[str] = []
    if winner.error:
        reasons.append(winner.error)
    if not winner.text:
        reasons.append("empty OCR output")
    if winner.source_letter_count < MIN_SOURCE_LETTERS:
        reasons.append("too few source-script letters")
    elif winner.source_letter_count < MIN_AUTOMATIC_LINE_LETTERS:
        reasons.append("unusually short OCR output")
    if winner.script_purity < MIN_SCRIPT_PURITY:
        reasons.append("low Unicode script purity")
    if winner.detected_script not in {winner.expected_script, "mixed"}:
        reasons.append("model/script disagreement")
    if winner.detected_script == "mixed":
        reasons.append("mixed-script output requires review")
    if winner.text_quality < MIN_TEXT_QUALITY:
        reasons.append("low linguistic text quality")
    if winner.confidence is None:
        reasons.append("OCR confidence unavailable")
    elif winner.confidence < MIN_OCR_CONFIDENCE:
        reasons.append("low OCR confidence")
    if len(ordered) > 1:
        runner_up = ordered[1]
        both_plausible = (
            winner.script_purity >= MIN_SCRIPT_PURITY
            and runner_up.script_purity >= MIN_SCRIPT_PURITY
            and winner.text_quality >= MIN_TEXT_QUALITY
            and runner_up.text_quality >= MIN_TEXT_QUALITY
            and winner.expected_script != runner_up.expected_script
        )
        if both_plausible:
            reasons.append("Hindi and Gurmukhi candidates both look plausible")
    accepted = not reasons
    return winner, accepted, reasons


def _order_quad(points: np.ndarray) -> np.ndarray:
    ordered = np.zeros((4, 2), dtype=np.float32)
    sums = points.sum(axis=1)
    differences = np.diff(points, axis=1).reshape(-1)
    ordered[0] = points[np.argmin(sums)]
    ordered[2] = points[np.argmax(sums)]
    ordered[1] = points[np.argmin(differences)]
    ordered[3] = points[np.argmax(differences)]
    return ordered


def _perspective_rectify(rgb: np.ndarray) -> tuple[np.ndarray, bool]:
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(blurred, 55, 160)
    edges = cv2.morphologyEx(
        edges, cv2.MORPH_CLOSE, np.ones((5, 5), np.uint8), iterations=2
    )
    contours, _ = cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    page_area = rgb.shape[0] * rgb.shape[1]
    for contour in sorted(contours, key=cv2.contourArea, reverse=True)[:12]:
        if cv2.contourArea(contour) < 0.32 * page_area:
            break
        perimeter = cv2.arcLength(contour, True)
        approximation = cv2.approxPolyDP(contour, 0.02 * perimeter, True)
        if len(approximation) != 4 or not cv2.isContourConvex(approximation):
            continue
        points = _order_quad(approximation.reshape(4, 2).astype(np.float32))
        top_width = np.linalg.norm(points[1] - points[0])
        bottom_width = np.linalg.norm(points[2] - points[3])
        left_height = np.linalg.norm(points[3] - points[0])
        right_height = np.linalg.norm(points[2] - points[1])
        width = int(max(top_width, bottom_width))
        height = int(max(left_height, right_height))
        if width < 300 or height < 300:
            continue
        destination = np.array(
            [[0, 0], [width - 1, 0], [width - 1, height - 1], [0, height - 1]],
            dtype=np.float32,
        )
        transform = cv2.getPerspectiveTransform(points, destination)
        return cv2.warpPerspective(rgb, transform, (width, height)), True
    return rgb, False


def _estimate_skew(gray: np.ndarray) -> float:
    edges = cv2.Canny(gray, 55, 150)
    lines = cv2.HoughLinesP(
        edges,
        1,
        np.pi / 180,
        threshold=max(40, gray.shape[1] // 12),
        minLineLength=max(80, gray.shape[1] // 8),
        maxLineGap=25,
    )
    if lines is None:
        return 0.0
    angles: list[float] = []
    for x1, y1, x2, y2 in lines[:, 0]:
        angle = math.degrees(math.atan2(y2 - y1, x2 - x1))
        if -12.0 <= angle <= 12.0:
            angles.append(angle)
    return float(statistics.median(angles)) if len(angles) >= 3 else 0.0


def _rotate_bound(rgb: np.ndarray, angle: float) -> np.ndarray:
    height, width = rgb.shape[:2]
    center = (width / 2.0, height / 2.0)
    matrix = cv2.getRotationMatrix2D(center, angle, 1.0)
    cosine, sine = abs(matrix[0, 0]), abs(matrix[0, 1])
    new_width = int(height * sine + width * cosine)
    new_height = int(height * cosine + width * sine)
    matrix[0, 2] += new_width / 2 - center[0]
    matrix[1, 2] += new_height / 2 - center[1]
    return cv2.warpAffine(
        rgb,
        matrix,
        (new_width, new_height),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_CONSTANT,
        borderValue=(255, 255, 255),
    )


def _optional_tesseract_orientation(rgb: np.ndarray) -> tuple[np.ndarray, bool]:
    try:
        pytesseract, tessdata_option = _prepare_tesseract(("osd",))

        result = pytesseract.image_to_osd(
            Image.fromarray(rgb),
            output_type=pytesseract.Output.DICT,
            config=tessdata_option,
        )
        rotation = int(result.get("rotate", 0) or 0)
        confidence = float(result.get("orientation_conf", 0.0) or 0.0)
        if rotation in {90, 180, 270} and confidence >= 8.0:
            return np.rot90(rgb, k={90: 3, 180: 2, 270: 1}[rotation]).copy(), True
    except Exception:
        pass
    return rgb, False


def correct_document(image: Image.Image) -> tuple[np.ndarray, list[str]]:
    """Apply conservative orientation, boundary, perspective, and skew correction."""

    operations = ["exif_orientation"]
    image = ImageOps.exif_transpose(image).convert("RGB")
    rgb = np.asarray(image)
    if max(rgb.shape[:2]) > MAX_IMAGE_SIDE:
        scale = MAX_IMAGE_SIDE / max(rgb.shape[:2])
        rgb = cv2.resize(
            rgb,
            (round(rgb.shape[1] * scale), round(rgb.shape[0] * scale)),
            interpolation=cv2.INTER_AREA,
        )
        operations.append("bounded_resize")
    rgb, oriented = _optional_tesseract_orientation(rgb)
    if oriented:
        operations.append("tesseract_orientation")
    rgb, rectified = _perspective_rectify(rgb)
    if rectified:
        operations.append("perspective_rectification")
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
    angle = _estimate_skew(gray)
    if 0.35 <= abs(angle) <= 10.0:
        rgb = _rotate_bound(rgb, angle)
        operations.append(f"deskew_{angle:.2f}_degrees")
    return rgb, operations


def enhance_page(corrected_rgb: np.ndarray, operations: list[str] | None = None) -> EnhancedPage:
    """Create mild recognition variants without replacing the preserved page."""

    applied = list(operations or [])
    gray = cv2.cvtColor(corrected_rgb, cv2.COLOR_RGB2GRAY)
    background_kernel = max(31, (min(gray.shape[:2]) // 18) | 1)
    # Dark ink on light paper must be removed from the *background estimate*.
    # A large closing operation fills narrow dark strokes while retaining slow
    # illumination/shadow changes. MORPH_OPEN would instead preserve the ink in
    # the divisor and can flatten faint handwriting to almost zero contrast.
    background = cv2.morphologyEx(
        gray,
        cv2.MORPH_CLOSE,
        cv2.getStructuringElement(
            cv2.MORPH_ELLIPSE, (background_kernel, background_kernel)
        ),
    )
    illumination = cv2.divide(gray, np.maximum(background, 1), scale=245)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(illumination)
    enhanced = cv2.bilateralFilter(enhanced, 5, 32, 32)
    blurred = cv2.GaussianBlur(enhanced, (0, 0), 1.0)
    enhanced = cv2.addWeighted(enhanced, 1.25, blurred, -0.25, 0)
    threshold = cv2.adaptiveThreshold(
        enhanced,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        41,
        13,
    )
    applied.extend(
        ["illumination_correction", "clahe", "mild_bilateral_denoise", "mild_unsharp"]
    )
    return EnhancedPage(
        corrected_rgb=corrected_rgb,
        enhanced_rgb=cv2.cvtColor(enhanced, cv2.COLOR_GRAY2RGB),
        enhanced_gray=enhanced,
        threshold=threshold,
        operations=applied,
    )


def _clip_bbox(bbox: Sequence[float | int], width: int, height: int) -> BBox:
    x1, y1, x2, y2 = (int(round(float(value))) for value in bbox)
    x1 = max(0, min(width - 1, x1))
    y1 = max(0, min(height - 1, y1))
    x2 = max(x1 + 1, min(width, x2))
    y2 = max(y1 + 1, min(height, y2))
    return x1, y1, x2, y2


def _valid_region(bbox: BBox, page_width: int, page_height: int) -> bool:
    x1, y1, x2, y2 = bbox
    width, height = x2 - x1, y2 - y1
    area = width * height
    if width < MIN_LINE_WIDTH or height < MIN_LINE_HEIGHT or area < MIN_LINE_AREA:
        return False
    border = max(3, round(min(page_width, page_height) * 0.006))
    touches_edge = x1 <= border or y1 <= border or x2 >= page_width - border or y2 >= page_height - border
    is_border_artifact = touches_edge and (
        height > 0.72 * page_height or width > 0.94 * page_width
    )
    return not is_border_artifact


def _vertical_overlap(first: BBox, second: BBox) -> float:
    overlap = max(0, min(first[3], second[3]) - max(first[1], second[1]))
    return overlap / max(1, min(first[3] - first[1], second[3] - second[1]))


def merge_text_regions(regions: Sequence[TextRegion]) -> list[TextRegion]:
    """Merge adjacent fragments sharing a handwriting baseline into logical lines."""

    pending = sorted(regions, key=lambda region: (region.bbox[1], region.bbox[0]))
    merged: list[TextRegion] = []
    for region in pending:
        x1, y1, x2, y2 = region.bbox
        match_index: int | None = None
        best_gap = float("inf")
        for index, existing in enumerate(merged):
            ex1, ey1, ex2, ey2 = existing.bbox
            mean_height = ((y2 - y1) + (ey2 - ey1)) / 2.0
            baseline_delta = abs(y2 - ey2)
            gap = max(0, max(x1, ex1) - min(x2, ex2))
            same_line = (
                _vertical_overlap(region.bbox, existing.bbox) >= 0.45
                or baseline_delta <= 0.42 * mean_height
            )
            if same_line and gap <= max(28.0, 2.2 * mean_height) and gap < best_gap:
                match_index, best_gap = index, gap
        if match_index is None:
            merged.append(region)
            continue
        existing = merged[match_index]
        confidence_values = [
            value
            for value in (existing.detection_confidence, region.detection_confidence)
            if value is not None
        ]
        merged[match_index] = TextRegion(
            bbox=(
                min(existing.bbox[0], x1),
                min(existing.bbox[1], y1),
                max(existing.bbox[2], x2),
                max(existing.bbox[3], y2),
            ),
            detection_confidence=(
                sum(confidence_values) / len(confidence_values)
                if confidence_values
                else None
            ),
            detector=existing.detector,
        )
    return merged


def sort_reading_order(
    regions: Sequence[TextRegion], page_width: int | None = None
) -> list[TextRegion]:
    """Sort single-column pages by rows and detect obvious two-column layouts."""

    if not regions:
        return []
    heights = [region.bbox[3] - region.bbox[1] for region in regions]
    median_height = float(statistics.median(heights))
    if page_width and len(regions) >= 6:
        centers = sorted((region.bbox[0] + region.bbox[2]) / 2 for region in regions)
        gaps = [(centers[index + 1] - centers[index], index) for index in range(len(centers) - 1)]
        largest_gap, split_index = max(gaps, default=(0.0, 0))
        if largest_gap > 0.22 * page_width and split_index >= 1 and len(centers) - split_index >= 3:
            boundary = (centers[split_index] + centers[split_index + 1]) / 2
            left = [region for region in regions if (region.bbox[0] + region.bbox[2]) / 2 <= boundary]
            right = [region for region in regions if region not in left]
            return sorted(left, key=lambda region: (region.bbox[1], region.bbox[0])) + sorted(
                right, key=lambda region: (region.bbox[1], region.bbox[0])
            )

    rows: list[list[TextRegion]] = []
    for region in sorted(regions, key=lambda item: (item.bbox[1], item.bbox[0])):
        center_y = (region.bbox[1] + region.bbox[3]) / 2
        target: list[TextRegion] | None = None
        for row in rows:
            row_center = statistics.mean((item.bbox[1] + item.bbox[3]) / 2 for item in row)
            if abs(center_y - row_center) <= 0.55 * median_height:
                target = row
                break
        if target is None:
            target = []
            rows.append(target)
        target.append(region)
    ordered: list[TextRegion] = []
    for row in sorted(rows, key=lambda value: min(item.bbox[1] for item in value)):
        ordered.extend(sorted(row, key=lambda item: item.bbox[0]))
    return ordered


def _opencv_detect(page: EnhancedPage) -> list[TextRegion]:
    ink = cv2.bitwise_not(page.threshold)
    height, width = ink.shape
    border_x, border_y = max(4, width // 120), max(4, height // 120)
    ink[:border_y, :] = 0
    ink[-border_y:, :] = 0
    ink[:, :border_x] = 0
    ink[:, -border_x:] = 0

    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(40, width // 7), 1))
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(40, height // 7)))
    rules = cv2.bitwise_or(
        cv2.morphologyEx(ink, cv2.MORPH_OPEN, horizontal_kernel),
        cv2.morphologyEx(ink, cv2.MORPH_OPEN, vertical_kernel),
    )
    text_ink = cv2.subtract(ink, rules)
    connect_width = max(18, width // 55)
    connected = cv2.morphologyEx(
        text_ink,
        cv2.MORPH_CLOSE,
        cv2.getStructuringElement(cv2.MORPH_RECT, (connect_width, 3)),
        iterations=1,
    )
    connected = cv2.dilate(
        connected,
        cv2.getStructuringElement(cv2.MORPH_RECT, (max(9, width // 100), 3)),
        iterations=1,
    )
    contours, _ = cv2.findContours(connected, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    regions: list[TextRegion] = []
    for contour in contours:
        x, y, box_width, box_height = cv2.boundingRect(contour)
        bbox = _clip_bbox((x, y, x + box_width, y + box_height), width, height)
        if _valid_region(bbox, width, height):
            regions.append(TextRegion(bbox, None, "opencv_projection_fallback"))

    if not regions:
        row_ink = np.count_nonzero(text_ink, axis=1)
        active = row_ink > max(3, int(width * 0.006))
        start: int | None = None
        for index, is_active in enumerate([*active.tolist(), False]):
            if is_active and start is None:
                start = index
            elif not is_active and start is not None:
                if index - start >= MIN_LINE_HEIGHT:
                    columns = np.where(np.any(text_ink[start:index] > 0, axis=0))[0]
                    if columns.size:
                        bbox = _clip_bbox(
                            (columns[0], start, columns[-1] + 1, index), width, height
                        )
                        if _valid_region(bbox, width, height):
                            regions.append(
                                TextRegion(bbox, None, "opencv_projection_fallback")
                            )
                start = None
    return sort_reading_order(merge_text_regions(regions), width)


class ModelRegistry:
    """Lazy, single-instance model registry optimized for Kaggle GPU memory."""

    def __init__(self) -> None:
        import torch

        self.torch = torch
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        if self.device.type == "cuda":
            major, _minor = torch.cuda.get_device_capability(0)
            self.dtype = torch.bfloat16 if major >= 8 else torch.float16
            self.gpu_name = torch.cuda.get_device_name(0)
        else:
            self.dtype = torch.float32
            self.gpu_name = "CPU"
        self.detector: Any | None = None
        self.hindi_processor: Any | None = None
        self.hindi_model: Any | None = None
        self.gurmukhi_processor: Any | None = None
        self.gurmukhi_model: Any | None = None
        self.translation_tokenizer: Any | None = None
        self.translation_model: Any | None = None
        self.status: dict[str, dict[str, Any]] = {
            "detector": {"id": "surya.detection.DetectionPredictor", "loaded": False},
            "hindi": {
                "id": HINDI_MODEL_ID,
                "primary_id": HINDI_MODEL_ID,
                "loaded": False,
                "fallback": False,
            },
            "gurmukhi": {
                "id": GURMUKHI_MODEL_ID,
                "primary_id": GURMUKHI_MODEL_ID,
                "loaded": False,
                "fallback": True,
                "fallback_reason": "no verified dedicated line-level Gurmukhi handwriting checkpoint",
            },
            "translation": {"id": TRANSLATION_MODEL_ID, "loaded": False},
        }

    def autocast(self) -> contextlib.AbstractContextManager[Any]:
        if self.device.type != "cuda":
            return contextlib.nullcontext()
        return self.torch.autocast(device_type="cuda", dtype=self.dtype)

    def _release(self, *names: str) -> None:
        for name in names:
            setattr(self, name, None)
        gc.collect()
        if self.device.type == "cuda":
            self.torch.cuda.empty_cache()

    def release_detector(self) -> None:
        self._release("detector")

    def release_recognizers(self) -> None:
        self._release(
            "hindi_processor", "hindi_model", "gurmukhi_processor", "gurmukhi_model"
        )

    def release_translation(self) -> None:
        self._release("translation_tokenizer", "translation_model")

    def load_detector(self) -> Any:
        if self.detector is None:
            # Surya is an optional Kaggle dependency. Loading by module name
            # keeps local editors usable when only the OpenCV fallback is
            # installed, while still raising a normal import error at runtime.
            detection_module = importlib.import_module("surya.detection")
            detection_predictor = getattr(detection_module, "DetectionPredictor")

            # Own the detector in this process so releasing this registry also
            # releases GPU memory; the default constructor starts a shared
            # server that can outlive a Kaggle notebook run.
            self.detector = detection_predictor.local(
                device=self.device,
                dtype=self.dtype,
            )
            self.status["detector"]["loaded"] = True
        return self.detector

    def load_hindi(self) -> tuple[Any, Any]:
        if self.hindi_model is None:
            from transformers import (
                RobertaTokenizer,
                TrOCRProcessor,
                ViTImageProcessorPil,
                VisionEncoderDecoderModel,
            )

            # This checkpoint predates the modern AutoImageProcessor metadata.
            # Construct its documented ViT + Roberta processor explicitly so it
            # remains loadable on Transformers 5 without changing the model.
            image_processor = _from_pretrained_cached_first(
                ViTImageProcessorPil, HINDI_MODEL_ID
            )
            tokenizer = _from_pretrained_cached_first(RobertaTokenizer, HINDI_MODEL_ID)
            self.hindi_processor = TrOCRProcessor(
                image_processor=image_processor,
                tokenizer=tokenizer,
            )
            hindi_load_options: dict[str, Any] = {"low_cpu_mem_usage": False}
            if self.device.type == "cuda":
                hindi_load_options["dtype"] = self.dtype
            loaded_hindi = _from_pretrained_cached_first(
                VisionEncoderDecoderModel,
                HINDI_MODEL_ID,
                **hindi_load_options,
            )
            self.hindi_model = _finish_pretrained_model_load(
                loaded_hindi,
                self.device,
            )
            _use_max_new_tokens_only(self.hindi_model)
            self.hindi_model.eval()
            self.status["hindi"]["loaded"] = True
        return self.hindi_processor, self.hindi_model

    def load_gurmukhi(self) -> tuple[Any, Any]:
        if self.gurmukhi_model is None:
            from transformers import AutoModelForMultimodalLM, AutoProcessor

            self.gurmukhi_processor = _from_pretrained_cached_first(
                AutoProcessor, GURMUKHI_MODEL_ID
            )
            tokenizer = getattr(self.gurmukhi_processor, "tokenizer", None)
            if tokenizer is not None:
                tokenizer.padding_side = "left"
                if tokenizer.pad_token_id is None and tokenizer.eos_token_id is not None:
                    tokenizer.pad_token_id = tokenizer.eos_token_id
            loaded_gurmukhi = _from_pretrained_cached_first(
                AutoModelForMultimodalLM,
                GURMUKHI_MODEL_ID,
                # The checkpoint declares BF16.  Older AVX2-only CPUs emulate
                # BF16 very slowly, so explicitly load FP32 on CPU.
                dtype=self.dtype,
                low_cpu_mem_usage=False,
            )
            self.gurmukhi_model = _finish_pretrained_model_load(
                loaded_gurmukhi,
                self.device,
            )
            if tokenizer is not None:
                # The repository's nested text config contains a stale Qwen
                # token id, while its tokenizer/generation config use 2 and 0.
                # Set generation ids from the loaded tokenizer to guarantee
                # early stopping instead of decoding every token budget.
                self.gurmukhi_model.generation_config.eos_token_id = tokenizer.eos_token_id
                self.gurmukhi_model.generation_config.pad_token_id = tokenizer.pad_token_id
            _use_max_new_tokens_only(self.gurmukhi_model)
            self.gurmukhi_model.eval()
            self.status["gurmukhi"]["loaded"] = True
        return self.gurmukhi_processor, self.gurmukhi_model

    def load_translation(self) -> tuple[Any, Any]:
        if self.translation_model is None:
            from transformers import AutoModelForSeq2SeqLM, AutoTokenizer

            self.translation_tokenizer = _from_pretrained_cached_first(
                AutoTokenizer, TRANSLATION_MODEL_ID
            )
            translation_load_options: dict[str, Any] = {"low_cpu_mem_usage": False}
            if self.device.type == "cuda":
                translation_load_options["dtype"] = self.dtype
            loaded_translation = _from_pretrained_cached_first(
                AutoModelForSeq2SeqLM,
                TRANSLATION_MODEL_ID,
                **translation_load_options,
            )
            self.translation_model = _finish_pretrained_model_load(
                loaded_translation,
                self.device,
            )
            _use_max_new_tokens_only(self.translation_model)
            self.translation_model.eval()
            self.status["translation"]["loaded"] = True
        return self.translation_tokenizer, self.translation_model


def load_models() -> ModelRegistry:
    """Create a lazy registry; downloads occur once when each model is first used."""

    return ModelRegistry()


def _object_value(value: Any, key: str, default: Any = None) -> Any:
    if isinstance(value, dict):
        return value.get(key, default)
    return getattr(value, key, default)


def detect_text_regions(page: EnhancedPage, models: ModelRegistry) -> list[TextRegion]:
    """Use pretrained Surya line detection, with a robust OpenCV fallback."""

    height, width = page.enhanced_gray.shape
    detector: Any | None = None
    try:
        detector = models.load_detector()
        prediction = detector([Image.fromarray(page.enhanced_rgb)])[0]
        regions: list[TextRegion] = []
        for item in _object_value(prediction, "bboxes", []) or []:
            raw_bbox = _object_value(item, "bbox")
            if raw_bbox is None:
                polygon = _object_value(item, "polygon", [])
                if polygon:
                    points = np.asarray(polygon, dtype=np.float32)
                    raw_bbox = (
                        points[:, 0].min(),
                        points[:, 1].min(),
                        points[:, 0].max(),
                        points[:, 1].max(),
                    )
            if raw_bbox is None:
                continue
            bbox = _clip_bbox(raw_bbox, width, height)
            if not _valid_region(bbox, width, height):
                continue
            confidence_value = _object_value(item, "confidence")
            confidence = (
                float(confidence_value) if confidence_value is not None else None
            )
            regions.append(TextRegion(bbox, confidence, "surya_text_detector"))
        regions = sort_reading_order(merge_text_regions(regions), width)
        if regions:
            return regions
        print("WARNING: Surya detected no valid lines; using the OpenCV fallback.")
    except Exception as exc:
        models.status["detector"]["error"] = f"{type(exc).__name__}: {exc}"
        print(f"WARNING: Surya line detection unavailable ({exc}); using OpenCV.")
    finally:
        detector = None
        models.release_detector()
    return _opencv_detect(page)


def prepare_crop(page_rgb: np.ndarray, bbox: BBox, padding: int = LINE_PADDING_PIXELS) -> Image.Image:
    """Return a padded RGB line crop without modifying page coordinates."""

    height, width = page_rgb.shape[:2]
    x1, y1, x2, y2 = bbox
    x1, y1 = max(0, x1 - padding), max(0, y1 - padding)
    x2, y2 = min(width, x2 + padding), min(height, y2 + padding)
    crop = Image.fromarray(page_rgb[y1:y2, x1:x2]).convert("RGB")
    return ImageOps.expand(crop, border=max(4, padding // 2), fill="white")


def _square_pad(image: Image.Image, padding: int = 8) -> Image.Image:
    """Center an OCR crop on a white square without stretching its aspect ratio."""

    source = image.convert("RGB")
    side = max(source.width, source.height) + 2 * max(0, padding)
    canvas = Image.new("RGB", (side, side), "white")
    canvas.paste(source, ((side - source.width) // 2, (side - source.height) // 2))
    return canvas


def _split_words(line: Image.Image) -> list[Image.Image]:
    """Split a line for the word-trained Hindi checkpoint when gaps are clear."""

    gray = np.asarray(line.convert("L"))
    threshold = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 31, 13
    )
    projection = np.count_nonzero(threshold, axis=0)
    active = projection > max(1, int(gray.shape[0] * 0.025))
    runs: list[tuple[int, int]] = []
    start: int | None = None
    for index, value in enumerate([*active.tolist(), False]):
        if value and start is None:
            start = index
        elif not value and start is not None:
            runs.append((start, index))
            start = None
    if len(runs) < 2:
        return [line]
    gap_limit = max(7, round(gray.shape[0] * 0.16))
    words: list[tuple[int, int]] = []
    current_start, current_end = runs[0]
    for run_start, run_end in runs[1:]:
        if run_start - current_end >= gap_limit:
            words.append((current_start, current_end))
            current_start, current_end = run_start, run_end
        else:
            current_end = run_end
    words.append((current_start, current_end))
    if not 2 <= len(words) <= 18:
        return [line]
    result = [
        ImageOps.expand(
            line.crop((max(0, x1 - 4), 0, min(line.width, x2 + 4), line.height)),
            border=4,
            fill="white",
        )
        for x1, x2 in words
        if x2 - x1 >= 4
    ]
    return result or [line]


def _batch(items: Sequence[Any], size: int) -> Iterator[Sequence[Any]]:
    for start in range(0, len(items), max(1, size)):
        yield items[start : start + max(1, size)]


def _report_progress(callback: ProgressCallback | None, message: str) -> None:
    if callback is not None:
        callback(message)


def _generation_confidences(model: Any, generated: Any) -> list[float | None]:
    try:
        transition = model.compute_transition_scores(
            generated.sequences,
            generated.scores,
            getattr(generated, "beam_indices", None),
            normalize_logits=True,
        )
        probabilities = transition.float().exp()
        token_ids = generated.sequences[:, -transition.shape[1] :]
        pad_token_id = getattr(model.generation_config, "pad_token_id", None)
        valid_tokens = (
            token_ids.ne(pad_token_id) if pad_token_id is not None else token_ids.ge(0)
        )
        confidences: list[float | None] = []
        for row, valid in zip(probabilities, valid_tokens, strict=True):
            selected = row[valid]
            confidences.append(float(selected.mean().item()) if selected.numel() else None)
        return confidences
    except Exception:
        return [None] * len(generated.sequences)


def _tesseract_candidate(crop: Image.Image, language: str, script: ScriptName) -> RecognitionCandidate:
    try:
        pytesseract, tessdata_option = _prepare_tesseract((language,))
        recognition_crop = crop.convert("RGB")
        if recognition_crop.height < 54:
            scale = min(3.0, max(1.5, 54.0 / max(1, recognition_crop.height)))
            recognition_crop = recognition_crop.resize(
                (
                    max(1, round(recognition_crop.width * scale)),
                    max(1, round(recognition_crop.height * scale)),
                ),
                Image.Resampling.LANCZOS,
            )
        config = " ".join(
            part for part in (tessdata_option, "--oem 1 --psm 7") if part
        )

        data = pytesseract.image_to_data(
            recognition_crop,
            lang=language,
            config=config,
            output_type=pytesseract.Output.DICT,
        )
        words: list[str] = []
        confidences: list[float] = []
        for text, confidence in zip(data.get("text", []), data.get("conf", []), strict=False):
            normalized = normalize_text(text)
            try:
                value = float(confidence)
            except (TypeError, ValueError):
                value = -1.0
            if normalized:
                words.append(normalized)
                if value >= 0:
                    confidences.append(value / 100.0)
        return RecognitionCandidate(
            expected_script=script,
            text=" ".join(words),
            confidence=(sum(confidences) / len(confidences) if confidences else None),
            model_id=f"tesseract:{language}",
            provider_kind="tesseract_fallback",
        )
    except Exception as exc:
        return RecognitionCandidate(
            expected_script=script,
            text="",
            confidence=None,
            model_id=f"tesseract:{language}",
            provider_kind="unavailable",
            error=f"Tesseract {language} unavailable: {exc}",
        )


def recognize_hindi(
    crops: Sequence[Image.Image],
    models: ModelRegistry,
    progress_callback: ProgressCallback | None = None,
) -> list[RecognitionCandidate]:
    """Recognize Devanagari with the required pretrained TrOCR checkpoint."""

    try:
        processor, model = models.load_hindi()
    except Exception as exc:
        models._release("hindi_processor", "hindi_model")
        models.status["hindi"].update(
            {
                "id": "tesseract:hin",
                "loaded": False,
                "fallback": True,
                "error": f"{type(exc).__name__}: {exc}",
            }
        )
        print(f"WARNING: Hindi TrOCR unavailable ({exc}); using Tesseract hin.")
        return [_tesseract_candidate(crop, "hin", "devanagari") for crop in crops]

    units: list[Image.Image] = []
    mapping: list[list[int]] = []
    for crop in crops:
        indices: list[int] = []
        for word in _split_words(crop):
            indices.append(len(units))
            units.append(_square_pad(word))
        mapping.append(indices)
    texts: list[str] = []
    confidences: list[float | None] = []
    torch = models.torch
    try:
        total_batches = max(1, math.ceil(len(units) / OCR_BATCH_SIZE))
        for batch_index, current in enumerate(_batch(units, OCR_BATCH_SIZE), start=1):
            _report_progress(
                progress_callback,
                f"Hindi handwriting OCR batch {batch_index}/{total_batches}",
            )
            pixels = processor(images=list(current), return_tensors="pt").pixel_values.to(
                models.device
            )
            if models.device.type == "cuda":
                pixels = pixels.to(dtype=models.dtype)
            with torch.inference_mode(), models.autocast():
                generated = model.generate(
                    pixel_values=pixels,
                    num_beams=4,
                    max_new_tokens=32,
                    return_dict_in_generate=True,
                    output_scores=True,
                )
            texts.extend(
                normalize_text(value)
                for value in processor.batch_decode(
                    generated.sequences, skip_special_tokens=True
                )
            )
            confidences.extend(_generation_confidences(model, generated))
    except Exception as exc:
        models.status["hindi"].update(
            {
                "id": "tesseract:hin",
                "fallback": True,
                "error": f"{type(exc).__name__}: {exc}",
            }
        )
        print(f"WARNING: Hindi TrOCR inference failed ({exc}); using Tesseract hin.")
        return [_tesseract_candidate(crop, "hin", "devanagari") for crop in crops]
    finally:
        processor = None
        model = None
        models._release("hindi_processor", "hindi_model")

    results: list[RecognitionCandidate] = []
    for indices in mapping:
        line_text = " ".join(texts[index] for index in indices if texts[index])
        values = [confidences[index] for index in indices if confidences[index] is not None]
        results.append(
            RecognitionCandidate(
                expected_script="devanagari",
                text=line_text,
                confidence=(sum(values) / len(values) if values else None),
                model_id=HINDI_MODEL_ID,
                provider_kind="dedicated_hindi_handwriting",
            )
        )
    return results


def _surya_prompts(processor: Any, count: int) -> list[str]:
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "image"},
                {
                    "type": "text",
                    # This is Surya 2's documented block-OCR training-time prompt.
                    # Unicode/source-quality gates below enforce the Gurmukhi route.
                    "text": "OCR this block image to HTML.",
                },
            ],
        }
    ]
    prompt = processor.apply_chat_template(
        messages, tokenize=False, add_generation_prompt=True
    )
    return [prompt] * count


def recognize_gurmukhi(
    crops: Sequence[Image.Image],
    models: ModelRegistry,
    progress_callback: ProgressCallback | None = None,
    *,
    fast: bool = False,
) -> list[RecognitionCandidate]:
    """Recognize Gurmukhi with the verified multilingual Surya fallback."""

    try:
        processor, model = models.load_gurmukhi()
    except Exception as exc:
        models._release("gurmukhi_processor", "gurmukhi_model")
        models.status["gurmukhi"].update(
            {
                "loaded": False,
                "id": "tesseract:pan",
                "fallback": True,
                "error": f"{type(exc).__name__}: {exc}",
            }
        )
        print(f"WARNING: Surya multilingual OCR unavailable ({exc}); using Tesseract pan.")
        return [_tesseract_candidate(crop, "pan", "gurmukhi") for crop in crops]

    results: list[RecognitionCandidate] = []
    torch = models.torch
    try:
        batch_size = VLM_BATCH_SIZE if models.device.type == "cuda" else (2 if fast else 1)
        token_budget = 128 if models.device.type == "cuda" else (48 if fast else 64)
        total_batches = max(1, math.ceil(len(crops) / batch_size))
        for batch_index, current in enumerate(_batch(crops, batch_size), start=1):
            _report_progress(
                progress_callback,
                (
                    f"Punjabi/Gurmukhi OCR batch {batch_index}/{total_batches} "
                    f"on {models.gpu_name}"
                ),
            )
            prompts = _surya_prompts(processor, len(current))
            inputs = processor(
                text=prompts,
                images=list(current),
                padding=True,
                return_tensors="pt",
            )
            inputs = {
                key: value.to(models.device) if hasattr(value, "to") else value
                for key, value in inputs.items()
            }
            with torch.inference_mode(), models.autocast():
                generated = model.generate(
                    **inputs,
                    max_new_tokens=token_budget,
                    do_sample=False,
                    return_dict_in_generate=True,
                    output_scores=True,
                )
            input_length = inputs["input_ids"].shape[1]
            new_tokens = generated.sequences[:, input_length:]
            decoded = processor.batch_decode(new_tokens, skip_special_tokens=True)
            batch_confidences = _generation_confidences(model, generated)
            for text, confidence in zip(decoded, batch_confidences, strict=True):
                results.append(
                    RecognitionCandidate(
                        expected_script="gurmukhi",
                        text=normalize_text(text),
                        confidence=confidence,
                        model_id=GURMUKHI_MODEL_ID,
                        provider_kind="multilingual_handwriting_fallback",
                    )
                )
    except Exception as exc:
        models.status["gurmukhi"].update(
            {
                "id": "tesseract:pan",
                "fallback": True,
                "error": f"{type(exc).__name__}: {exc}",
            }
        )
        print(f"WARNING: Surya Gurmukhi inference failed ({exc}); using Tesseract pan.")
        results = [_tesseract_candidate(crop, "pan", "gurmukhi") for crop in crops]
    processor = None
    model = None
    models._release("gurmukhi_processor", "gurmukhi_model")
    return results


def _candidate_for_audit(candidate: RecognitionCandidate) -> dict[str, Any]:
    return {
        "expected_script": candidate.expected_script,
        "detected_script": candidate.detected_script,
        "text": candidate.text,
        "confidence": candidate.confidence,
        "script_purity": candidate.script_purity,
        "text_quality": candidate.text_quality,
        "source_letter_count": candidate.source_letter_count,
        "score": candidate.score,
        "model_id": candidate.model_id,
        "provider_kind": candidate.provider_kind,
        "error": candidate.error,
    }


def _language_code(script: ScriptName) -> str | None:
    return {"gurmukhi": "pan_Guru", "devanagari": "hin_Deva"}.get(script)


def _can_share_translation_context(
    previous: LineResult,
    current: LineResult,
    group_size: int,
    group_characters: int,
) -> bool:
    """Return whether two validated scan lines can safely form one source unit."""

    if group_size >= TRANSLATION_CONTEXT_MAX_LINES:
        return False
    if group_characters + len(current.text) + 1 > TRANSLATION_CONTEXT_MAX_CHARACTERS:
        return False
    if previous.page_number != current.page_number or previous.script != current.script:
        return False
    if current.line != previous.line + 1:
        return False
    # Native text already arrives as reliable structural units. Avoid merging
    # separate Word/PDF objects because those mappings must remain exact.
    native_kinds = {"docx_native", "pymupdf_native_text"}
    if previous.provider_kind in native_kinds or current.provider_kind in native_kinds:
        return False
    if re.search(r"[.!?।॥]\s*$", previous.text):
        return False
    previous_height = max(1, previous.bbox[3] - previous.bbox[1])
    current_height = max(1, current.bbox[3] - current.bbox[1])
    vertical_gap = current.bbox[1] - previous.bbox[3]
    return -0.40 * max(previous_height, current_height) <= vertical_gap <= 1.80 * max(
        previous_height,
        current_height,
    )


def build_translation_units(lines: Sequence[LineResult]) -> list[TranslationUnit]:
    """Group adjacent validated scan lines while keeping rejected lines as barriers."""

    units: list[TranslationUnit] = []
    current: list[LineResult] = []

    def flush() -> None:
        nonlocal current
        if not current:
            return
        source = normalize_translation_source(" ".join(line.text for line in current))
        group_id = f"p{current[0].page_number}:l{current[0].line}-{current[-1].line}"
        units.append(
            TranslationUnit(
                lines=current,
                script=current[0].script,
                source=source,
                protected=protect_translation_source(source),
                group_id=group_id,
            )
        )
        current = []

    for line in lines:
        eligible = line.accepted and _language_code(line.script) is not None
        if not eligible:
            flush()
            continue
        if current and not _can_share_translation_context(
            current[-1],
            line,
            len(current),
            sum(len(item.text) for item in current),
        ):
            flush()
        current.append(line)
    flush()
    return units


def _split_context_translation(
    translation: str,
    source_lines: Sequence[LineResult],
    protected_tokens: Sequence[str] = (),
) -> list[str] | None:
    """Split one contextual translation without dropping or duplicating words."""

    if len(source_lines) == 1:
        return [translation]
    atomic = translation
    atomic_values: dict[str, str] = {}
    for index, token in enumerate(sorted(set(protected_tokens), key=len, reverse=True)):
        if token not in atomic:
            continue
        placeholder = f"ZXQATOMIC{index}QXZ"
        atomic = atomic.replace(token, placeholder)
        atomic_values[placeholder] = token
    words = atomic.split()
    if len(words) < len(source_lines):
        return None
    weights = [
        max(1, int(script_statistics(line.text)["letters"])) for line in source_lines
    ]
    remaining_words = len(words)
    remaining_weight = sum(weights)
    values: list[str] = []
    start = 0
    for index, weight in enumerate(weights[:-1]):
        remaining_lines = len(weights) - index - 1
        size = round(remaining_words * weight / max(1, remaining_weight))
        size = max(1, min(size, remaining_words - remaining_lines))
        values.append(" ".join(words[start : start + size]))
        start += size
        remaining_words -= size
        remaining_weight -= weight
    values.append(" ".join(words[start:]))
    for index, value in enumerate(values):
        for placeholder, token in atomic_values.items():
            value = value.replace(placeholder, token)
        values[index] = value
    return values if all(values) else None


def translate_text(
    lines: list[LineResult],
    models: ModelRegistry,
    progress_callback: ProgressCallback | None = None,
    num_beams: int = 4,
) -> float:
    """Translate validated source units with context, exact tokens, and QA."""

    started = time.perf_counter()
    units = build_translation_units(lines)
    if not units:
        return time.perf_counter() - started
    try:
        tokenizer, model = models.load_translation()
    except Exception as exc:
        models.release_translation()
        for unit in units:
            for line in unit.lines:
                line.translation_status = "model_unavailable"
                line.translation_error = f"{type(exc).__name__}: {exc}"
                line.translation_review_reasons = ["translation model unavailable"]
                line.review_required = True
        models.status["translation"]["error"] = f"{type(exc).__name__}: {exc}"
        return time.perf_counter() - started

    torch = models.torch
    target_id = tokenizer.convert_tokens_to_ids("eng_Latn")
    requested_beams = max(4, int(num_beams))

    def translate_batch(
        current: Sequence[TranslationUnit],
        beams: int,
    ) -> list[tuple[str, float | None]]:
        encoded = tokenizer(
            [unit.protected.text for unit in current],
            padding=True,
            truncation=True,
            max_length=384,
            return_tensors="pt",
        )
        encoded = {
            key: value.to(models.device) if hasattr(value, "to") else value
            for key, value in encoded.items()
        }
        with torch.inference_mode(), models.autocast():
            output = model.generate(
                **encoded,
                forced_bos_token_id=target_id,
                max_new_tokens=384,
                num_beams=beams,
                length_penalty=1.0,
                repetition_penalty=1.05,
                no_repeat_ngram_size=3,
                early_stopping=True,
                return_dict_in_generate=True,
                output_scores=True,
            )
        decoded = [
            normalize_text(value)
            for value in tokenizer.batch_decode(
                output.sequences,
                skip_special_tokens=True,
            )
        ]
        confidences = _generation_confidences(model, output)
        return list(zip(decoded, confidences, strict=True))

    def apply_translation(
        unit: TranslationUnit,
        raw_translation: str,
        confidence: float | None,
    ) -> list[str]:
        restored, missing_tokens = restore_translation_source(
            raw_translation,
            unit.protected,
        )
        quality, reasons = translation_quality(
            unit.source,
            restored,
            missing_tokens,
        )
        if quality < 0.55 and not reasons:
            reasons.append("translation quality score below acceptance threshold")
        pieces = _split_context_translation(
            restored,
            unit.lines,
            unit.protected.source_tokens,
        )
        if pieces is None:
            reasons.append("context translation could not be mapped back to source lines")
        if reasons:
            return reasons
        assert pieces is not None
        for line, translation in zip(unit.lines, pieces, strict=True):
            line.english = translation
            line.translation_confidence = confidence
            line.translation_quality = quality
            line.translation_review_reasons = []
            line.translation_context_group = unit.group_id
            if translation:
                line.translation_status = "translated"
                # This flag may already be false after source validation. Keep
                # unrelated OCR review reasons intact if a caller supplied any.
                line.review_required = bool(line.review_reasons)
            else:
                line.translation_status = "empty_translation"
                line.translation_error = "translation model returned empty text"
                line.review_required = True
        return []

    def mark_translation_failure(
        unit: TranslationUnit,
        reasons: Sequence[str],
        exc: Exception | None = None,
    ) -> None:
        detail = f"{type(exc).__name__}: {exc}" if exc else "; ".join(reasons)
        for line in unit.lines:
            line.translation_status = "failed_validation" if reasons else "failed"
            line.translation_error = detail
            line.translation_review_reasons = list(reasons)
            line.review_required = True

    def singleton_units(unit: TranslationUnit) -> list[TranslationUnit]:
        return [
            TranslationUnit(
                lines=[line],
                script=line.script,
                source=normalize_translation_source(line.text),
                protected=protect_translation_source(normalize_translation_source(line.text)),
                group_id=f"p{line.page_number}:l{line.line}",
            )
            for line in unit.lines
        ]

    def retry_unit(unit: TranslationUnit) -> None:
        retry_beams = max(6, requested_beams + 1)
        retry_units = singleton_units(unit) if len(unit.lines) > 1 else [unit]
        for retry in retry_units:
            try:
                raw, confidence = translate_batch([retry], retry_beams)[0]
                reasons = apply_translation(retry, raw, confidence)
                if reasons:
                    mark_translation_failure(retry, reasons)
            except Exception as exc:
                mark_translation_failure(retry, [], exc)

    for script in ("gurmukhi", "devanagari"):
        group = [unit for unit in units if unit.script == script]
        source_language = _language_code(script)  # type: ignore[arg-type]
        if not group or source_language is None:
            continue
        tokenizer.src_lang = source_language
        total_batches = max(1, math.ceil(len(group) / TRANSLATION_BATCH_SIZE))
        for batch_index, current in enumerate(
            _batch(group, TRANSLATION_BATCH_SIZE), start=1
        ):
            _report_progress(
                progress_callback,
                f"English translation batch {batch_index}/{total_batches} ({source_language})",
            )
            try:
                translated = translate_batch(current, requested_beams)
                for unit, (raw, confidence) in zip(current, translated, strict=True):
                    reasons = apply_translation(unit, raw, confidence)
                    if reasons:
                        retry_unit(unit)
            except Exception as batch_exc:
                # One unusually shaped unit must not discard the full batch.
                for unit in current:
                    retry_unit(unit)
                models.status["translation"]["last_batch_error"] = (
                    f"{type(batch_exc).__name__}: {batch_exc}"
                )
    models.status["translation"].update(
        {
            "decoding": "quality_beam_search",
            "num_beams": requested_beams,
            "context_max_lines": TRANSLATION_CONTEXT_MAX_LINES,
            "protected_tokens": True,
            "output_validation": True,
        }
    )
    tokenizer = None
    model = None
    models.release_translation()
    return time.perf_counter() - started


def _draw_detected_lines(page_rgb: np.ndarray, lines: Sequence[LineResult]) -> np.ndarray:
    output = cv2.cvtColor(page_rgb.copy(), cv2.COLOR_RGB2BGR)
    for line in lines:
        x1, y1, x2, y2 = line.bbox
        color = (52, 180, 66) if line.accepted and not line.review_required else (25, 120, 245)
        cv2.rectangle(output, (x1, y1), (x2, y2), color, 2)
        label = str(line.line)
        cv2.rectangle(output, (x1, max(0, y1 - 22)), (x1 + 30, y1), color, -1)
        cv2.putText(
            output,
            label,
            (x1 + 4, max(14, y1 - 5)),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.5,
            (255, 255, 255),
            1,
            cv2.LINE_AA,
        )
    return output


def _line_transcription(line: LineResult) -> str:
    if line.accepted:
        return line.text
    return f"[REVIEW REQUIRED — line {line.line}] {line.text}".rstrip()


def _line_translation(line: LineResult) -> str:
    if line.translation_status in {
        "translated",
        "preserved_english",
        "preserved_structured",
    }:
        return line.english
    if line.review_required:
        return f"[REVIEW REQUIRED — line {line.line}]"
    return ""


def _word_document_line(line: LineResult) -> str:
    """Return one nonempty, provenance-safe Word paragraph per detected line."""

    english = normalize_text(line.english)
    source = normalize_text(line.text)
    if line.translation_status in {
        "translated",
        "preserved_english",
        "preserved_structured",
    } and english:
        return english
    if line.accepted and source:
        return (
            f"[Translation unavailable for line {line.line}; source retained] "
            f"{source}"
        )
    if source:
        return f"[Review required for line {line.line}; uncertain source retained] {source}"
    return f"[Review required; no reliable transcription for detected line {line.line}]"


def build_translated_docx(page: EnhancedPage, lines: Sequence[LineResult]) -> bytes:
    """Build a clean Word document containing every detected line in reading order."""

    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    source_height, source_width = page.corrected_rgb.shape[:2]
    page_width_inches = 11.69 if source_width >= source_height else 8.27
    page_height_inches = page_width_inches * source_height / max(1, source_width)
    page_height_inches = max(5.0, min(22.0, page_height_inches))
    section.page_width = Inches(page_width_inches)
    section.page_height = Inches(page_height_inches)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala UI")

    usable_height_points = max(1.0, (page_height_inches - 0.9) * 72.0)
    previous_bottom = 0
    previous_page = lines[0].page_number if lines else 1
    expected_paragraphs: list[str] = []
    for line in lines:
        if line.page_number != previous_page:
            document.add_page_break()
            previous_bottom = 0
            previous_page = line.page_number
        value = _word_document_line(line)
        expected_paragraphs.append(value)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(0)
        y_gap = max(0, line.bbox[1] - previous_bottom)
        paragraph.paragraph_format.space_before = Pt(
            min(18.0, y_gap / max(1, source_height) * usable_height_points)
        )
        previous_bottom = max(previous_bottom, line.bbox[3])
        run = paragraph.add_run(value)
        run.font.name = "Aptos" if line.translation_status == "translated" else "Nirmala UI"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala UI")
        source_line_height = max(1, line.bbox[3] - line.bbox[1])
        run.font.size = Pt(max(8.0, min(13.0, source_line_height * 0.32)))

    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    reopened = Document(BytesIO(content))
    actual_paragraphs = [
        paragraph.text for paragraph in reopened.paragraphs if paragraph.text
    ]
    if actual_paragraphs != expected_paragraphs:
        raise OSError("Generated Word document failed its ordered-line validation")
    return content


def save_results(
    *,
    input_image: Path,
    page: EnhancedPage,
    lines: list[LineResult],
    models: ModelRegistry,
    timings: dict[str, float],
    output_directory: Path = OUTPUT_DIRECTORY,
) -> dict[str, Any]:
    """Write primary Word output plus image, text, JSON, and CSV artifacts."""

    output_directory.mkdir(parents=True, exist_ok=True)
    corrected_bgr = cv2.cvtColor(page.corrected_rgb, cv2.COLOR_RGB2BGR)
    corrected_path = output_directory / "corrected_page.jpg"
    detected_path = output_directory / "detected_lines.jpg"
    corrected_written = cv2.imwrite(str(corrected_path), corrected_bgr)
    detected_written = cv2.imwrite(
        str(detected_path),
        _draw_detected_lines(page.corrected_rgb, lines),
    )
    if not corrected_written or not detected_written:
        failed = [
            str(path)
            for path, written in (
                (corrected_path, corrected_written),
                (detected_path, detected_written),
            )
            if not written
        ]
        raise OSError(f"Could not write output image(s): {', '.join(failed)}")

    transcription = "\n".join(_line_transcription(line) for line in lines)
    english_translation = "\n".join(
        value for value in (_line_translation(line) for line in lines) if value
    )
    (output_directory / "transcription.txt").write_text(
        transcription + "\n", encoding="utf-8"
    )
    (output_directory / "translation_en.txt").write_text(
        english_translation + "\n", encoding="utf-8"
    )
    translated_docx_path = output_directory / "translated_en.docx"
    translated_docx_path.write_bytes(build_translated_docx(page, lines))

    payload = {
        "schema_version": "1.0",
        "input_image": str(input_image),
        "output_directory": str(output_directory),
        "inference_only": True,
        "models": models.status,
        "preprocessing_operations": page.operations,
        "lines": [asdict(line) for line in lines],
        "transcription": transcription,
        "english_translation": english_translation,
        "primary_output": translated_docx_path.name,
        "timings_seconds": timings,
        "accuracy_warning": (
            "Handwriting OCR is probabilistic. Review every uncertain line; this output "
            "is not a certified legal or medical transcription or translation."
        ),
    }
    (output_directory / "result.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    csv_fields = [
        "page_number",
        "line",
        "bbox",
        "script",
        "text",
        "script_purity",
        "confidence",
        "text_quality",
        "recognition_model",
        "provider_kind",
        "accepted",
        "review_required",
        "review_reasons",
        "english",
        "translation_confidence",
        "translation_quality",
        "translation_review_reasons",
        "translation_context_group",
        "translation_status",
        "translation_error",
        "crop_file",
    ]
    with (output_directory / "result.csv").open(
        "w", encoding="utf-8-sig", newline=""
    ) as stream:
        writer = csv.DictWriter(stream, fieldnames=csv_fields)
        writer.writeheader()
        for line in lines:
            row = asdict(line)
            row["bbox"] = json.dumps(row["bbox"])
            row["review_reasons"] = "; ".join(line.review_reasons)
            row["translation_review_reasons"] = "; ".join(
                line.translation_review_reasons
            )
            writer.writerow({key: row.get(key) for key in csv_fields})
    return payload


def _resolve_input(explicit: str | None) -> Path:
    if explicit:
        candidate = Path(explicit).expanduser()
        if candidate.is_file():
            return candidate.resolve()
        if "..." not in explicit:
            raise FileNotFoundError(f"PAGE_IMAGE does not exist: {candidate}")
        print("PAGE_IMAGE still contains '...'; auto-detecting an image under /kaggle/input.")
    input_root = Path("/kaggle/input")
    if not input_root.is_dir():
        raise FileNotFoundError(
            "Set PAGE_IMAGE to an existing JPG/JPEG/PNG file; /kaggle/input is unavailable."
        )
    candidates = sorted(
        path
        for path in input_root.rglob("*")
        if path.is_file() and path.suffix.casefold() in {".jpg", ".jpeg", ".png"}
    )
    if not candidates:
        raise FileNotFoundError("No JPG/JPEG/PNG image was found under /kaggle/input")
    return candidates[0]


def process_page(
    image_path: Path,
    output_directory: Path = OUTPUT_DIRECTORY,
    progress_callback: ProgressCallback | None = None,
    processing_mode: ProcessingMode = "fast_cpu",
) -> dict[str, Any]:
    """Run preprocessing, OCR validation, translation, and Word generation."""

    if processing_mode not in {"fast_cpu", "line_accurate"}:
        raise ValueError(f"Unsupported processing mode: {processing_mode}")

    total_started = time.perf_counter()
    models = load_models()
    _report_progress(progress_callback, "Correcting and enhancing the uploaded page")
    with Image.open(image_path) as source:
        corrected_rgb, operations = correct_document(source)
    page = enhance_page(corrected_rgb, operations)

    detection_started = time.perf_counter()
    _report_progress(progress_callback, "Detecting logical handwritten text lines")
    regions = detect_text_regions(page, models)
    detection_time = time.perf_counter() - detection_started
    if not regions:
        raise RuntimeError("No credible text-line regions were detected")

    crops = [prepare_crop(page.enhanced_rgb, region.bbox) for region in regions]
    output_directory.mkdir(parents=True, exist_ok=True)
    crop_directory = output_directory / "line_crops"
    crop_directory.mkdir(parents=True, exist_ok=True)
    for stale in crop_directory.glob("line_*.jpg"):
        stale.unlink()
    crop_names: list[str] = []
    for index, crop in enumerate(crops, start=1):
        relative = Path("line_crops") / f"line_{index:03d}.jpg"
        crop.save(output_directory / relative, quality=95)
        crop_names.append(relative.as_posix())

    ocr_started = time.perf_counter()
    fast_mode = processing_mode == "fast_cpu"
    _report_progress(
        progress_callback,
        (
            f"Recognizing all {len(crops)} detected lines"
            + (" in faster CPU batches" if fast_mode else " with accuracy routing")
        ),
    )
    gurmukhi_candidates = recognize_gurmukhi(
        crops,
        models,
        progress_callback,
        fast=fast_mode,
    )
    unresolved_indices = [
        index
        for index, candidate in enumerate(gurmukhi_candidates)
        if not choose_best_recognition([candidate])[1]
    ]
    hindi_candidates = [
        RecognitionCandidate(
            expected_script="devanagari",
            text="",
            confidence=None,
            model_id=HINDI_MODEL_ID,
            provider_kind="not_routed",
            error="Hindi OCR not needed after validated Gurmukhi recognition",
        )
        for _crop in crops
    ]
    if unresolved_indices and fast_mode:
        # A fast installed-language-pack check prevents the optional Hindi
        # handwriting checkpoint download from stalling a CPU-first run.
        _report_progress(
            progress_callback,
            f"Fast Hindi fallback for {len(unresolved_indices)} unresolved lines",
        )
        for index in unresolved_indices:
            hindi_candidates[index] = _tesseract_candidate(
                crops[index], "hin", "devanagari"
            )
        models.status["hindi"].update(
            {
                "id": "tesseract:hin",
                "fallback": True,
                "skip_reason": "fast CPU mode avoids a second large model load",
            }
        )
    elif unresolved_indices:
        _report_progress(
            progress_callback,
            (
                f"Loading Hindi handwriting OCR for {len(unresolved_indices)} "
                "unresolved lines"
            ),
        )
        unresolved_hindi = recognize_hindi(
            [crops[index] for index in unresolved_indices],
            models,
            progress_callback,
        )
        for index, candidate in zip(unresolved_indices, unresolved_hindi, strict=True):
            hindi_candidates[index] = candidate
    else:
        models.status["hindi"]["skipped"] = True
        models.status["hindi"]["skip_reason"] = (
            "all lines passed validated Gurmukhi source gates"
        )
    models.release_recognizers()
    lines: list[LineResult] = []
    for index, (region, hindi, gurmukhi) in enumerate(
        zip(regions, hindi_candidates, gurmukhi_candidates, strict=True), start=1
    ):
        candidates = [hindi, gurmukhi]
        winner, accepted, reasons = choose_best_recognition(candidates)
        script = winner.detected_script
        lines.append(
            LineResult(
                line=index,
                bbox=region.bbox,
                crop_file=crop_names[index - 1],
                detection_confidence=region.detection_confidence,
                detector=region.detector,
                script=script,
                text=winner.text,
                script_purity=winner.script_purity,
                confidence=winner.confidence,
                text_quality=winner.text_quality,
                recognition_model=winner.model_id,
                provider_kind=winner.provider_kind,
                accepted=accepted,
                review_required=not accepted,
                review_reasons=reasons,
                candidates=[
                    _candidate_for_audit(enrich_candidate(candidate))
                    for candidate in candidates
                ],
            )
        )
    ocr_time = time.perf_counter() - ocr_started
    _report_progress(progress_callback, "Translating validated source lines to English")
    translation_time = translate_text(
        lines,
        models,
        progress_callback,
        num_beams=4 if fast_mode else 6,
    )
    total_time = time.perf_counter() - total_started
    timings = {
        "detection": detection_time,
        "ocr": ocr_time,
        "translation": translation_time,
        "total": total_time,
    }
    models.status["processing"] = {"mode": processing_mode}
    payload = save_results(
        input_image=image_path,
        page=page,
        lines=lines,
        models=models,
        timings=timings,
        output_directory=output_directory,
    )
    accepted_count = sum(line.accepted for line in lines)
    review_count = sum(line.review_required for line in lines)
    print("\nINFERENCE SUMMARY")
    print(f"GPU              : {models.gpu_name}")
    print(f"Hindi model      : {models.status['hindi']['id']}")
    gurmukhi_id = str(models.status["gurmukhi"]["id"])
    if gurmukhi_id == GURMUKHI_MODEL_ID:
        gurmukhi_note = "multilingual handwriting fallback; not a dedicated Gurmukhi HTR checkpoint"
    else:
        gurmukhi_note = "last-resort printed OCR fallback; not handwriting-specialized"
    print(f"Gurmukhi model   : {gurmukhi_id} ({gurmukhi_note})")
    print(f"Translation      : {models.status['translation']['id']}")
    print(f"Detected regions : {len(lines)}")
    print(f"Accepted lines   : {accepted_count}")
    print(f"Review required  : {review_count}")
    print(f"OCR time         : {ocr_time:.2f}s")
    print(f"Translation time : {translation_time:.2f}s")
    print(f"Total time       : {total_time:.2f}s")
    print("\nPUNJABI/HINDI TRANSCRIPTION")
    print(payload["transcription"])
    print("\nENGLISH TRANSLATION")
    print(payload["english_translation"])
    print(f"\nResults written to: {output_directory}")
    return payload


def _parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Inference-only handwritten Hindi/Gurmukhi page OCR"
    )
    parser.add_argument("--image", type=str, help="Override PAGE_IMAGE")
    parser.add_argument("--output", type=Path, default=OUTPUT_DIRECTORY)
    parser.add_argument(
        "--mode",
        choices=("fast_cpu", "line_accurate"),
        default="fast_cpu",
        help="Use faster CPU batching or the slower dual-neural-model accuracy path",
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = _parse_args(argv)
    try:
        image_path = _resolve_input(args.image if args.image is not None else PAGE_IMAGE)
        process_page(image_path, args.output, processing_mode=args.mode)
        return 0
    except KeyboardInterrupt:
        print("Cancelled by user.", file=sys.stderr)
        return 130
    except Exception as exc:
        print(f"ERROR: {type(exc).__name__}: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
